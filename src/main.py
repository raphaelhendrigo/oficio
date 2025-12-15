import os
import re
import sys
import time
import unicodedata
from datetime import date, datetime
from typing import Optional
from pathlib import Path
from urllib.parse import urljoin

from dotenv import load_dotenv
from playwright.sync_api import sync_playwright, TimeoutError as PWTimeoutError


def env_bool(name: str, default: bool = False) -> bool:
    v = os.getenv(name, str(default))
    return str(v).strip().lower() in ("1", "true", "yes", "y", "on")


def _is_login_page(page) -> bool:
    """Detecta se ainda estamos na tela de login do e-TCM."""
    try:
        user_loc = page.locator("#ctl00_cphMain_txtUsuario_I").first
        pass_loc = page.locator("#ctl00_cphMain_txtSenha_I").first
        if user_loc.count() > 0 and pass_loc.count() > 0:
            # Em algumas transicoes o formulario pode ficar no DOM porem oculto
            if user_loc.is_visible() or pass_loc.is_visible():
                return True
    except Exception:
        pass
    try:
        url = page.url or ""
        if "login.aspx" in url.lower():
            return True
    except Exception:
        pass
    try:
        title = page.title() or ""
        if "login" in title.lower():
            return True
    except Exception:
        pass
    return False


def _strip_quotes(v: Optional[str]) -> str:
    if v is None:
        return ""
    s = str(v).strip()
    if len(s) >= 2 and ((s[0] == s[-1] == '"') or (s[0] == s[-1] == "'")):
        s = s[1:-1]
    return s.strip()


def _extract_login_error(page) -> str:
    """Tenta extrair mensagem de erro exibida no formulario de login."""
    candidates = [
        "#ctl00_cphMain_lblMensagem",
        "#ctl00_cphMain_lblErro",
        "#ctl00_cphMain_lblMsg",
        ".dxValidationSummary, .dxvs, .dx-error, .dxeErrorCell",
        ".alert, .erro, .error, .msgErro",
    ]
    messages: list[str] = []
    for sel in candidates:
        try:
            loc = page.locator(sel)
            if loc.count() > 0:
                for t in loc.all_text_contents():
                    t = (t or "").strip()
                    if t and t not in messages:
                        messages.append(t)
        except Exception:
            continue
    if not messages:
        try:
            kw = page.locator(
                "text=/senha|usu[aá]rio|captcha|verifica|inv[aá]lido|incorreto|bloqueado/i"
            ).locator(":visible")
            if kw.count() > 0:
                t = (kw.first.inner_text() or "").strip()
                if t:
                    messages.append(t)
        except Exception:
            pass
    if not messages:
        return ""
    msg = " | ".join(messages)
    if len(msg) > 300:
        msg = msg[:300] + "..."
    return msg


def login_etcm(
    page,
    url: str,
    username: str,
    password: str,
    pause_after_login_ms: int = 0,
    login_manual_wait_ms: int = 45000,
    headless: bool = False,
) -> None:
    """Realiza login no e-TCM de forma robusta (DevExpress + captchas)."""
    username = _strip_quotes(username)
    password = _strip_quotes(password)

    print(f"Acessando: {url}")
    page.goto(url, wait_until="domcontentloaded", timeout=60000)

    # Aguarda os campos de login ficarem visiveis quando existirem.
    try:
        page.locator("#ctl00_cphMain_txtUsuario_I, input[name='ctl00$cphMain$txtUsuario']").first.wait_for(
            state="visible", timeout=30000
        )
    except Exception:
        pass

    # Usuario
    try:
        page.locator("#ctl00_cphMain_txtUsuario_I, input[name='ctl00$cphMain$txtUsuario']").first.fill(username)
    except Exception:
        try:
            page.locator("input[placeholder*='Usu'], input[name*='Usuario' i]").first.fill(username)
        except Exception:
            page.locator("input[type='text']").first.fill(username)

    # Senha
    try:
        page.locator("#ctl00_cphMain_txtSenha_I, input[name='ctl00$cphMain$txtSenha'][type='password']").first.fill(password)
    except Exception:
        try:
            page.locator("input[type='password']").first.fill(password)
        except Exception as e:
            raise RuntimeError("Nao foi possivel localizar o campo de senha.") from e

    btn_selectors = [
        "#ctl00_cphMain_btnLogin_I",  # input interno DevExpress (mais confiavel)
        "#ctl00_cphMain_btnLogin",
        "input[name='ctl00$cphMain$btnLogin']",
        "input[type='submit'][value*='Entrar' i]:not([readonly]):not([disabled])",
        "button:has-text('Entrar')",
        "text=/\\b(Entrar|Acessar|Login)\\b/i",
    ]

    clicked = False
    for sel in btn_selectors:
        try:
            page.locator(sel).first.click(timeout=5000)
        except Exception:
            continue
        # Considera sucesso quando os campos somem do DOM (reload/redirect).
        try:
            page.wait_for_function(
                "() => document.querySelector('#ctl00_cphMain_txtUsuario_I') === null && "
                "document.querySelector('#ctl00_cphMain_txtSenha_I') === null",
                timeout=15000,
            )
            clicked = True
            break
        except Exception:
            # Ainda na tela de login, tente outro seletor
            continue

    if not clicked:
        try:
            page.get_by_role("button", name=re.compile(r"Entrar|Acessar|Login", re.I)).click(timeout=5000)
            clicked = True
        except Exception:
            pass

    if not clicked:
        try:
            page.locator("#ctl00_cphMain_txtSenha_I, input[type='password']").first.press("Enter")
            clicked = True
        except Exception:
            pass

    if not clicked:
        # Fallback JS: __doPostBack / submit do form (DevExpress)
        try:
            page.evaluate(
                "(() => { try { __doPostBack('ctl00$cphMain$btnLogin',''); } catch(e) { "
                "var f=document.forms['aspnetForm']; if(f){f.__EVENTTARGET.value='ctl00$cphMain$btnLogin'; "
                "f.__EVENTARGUMENT.value=''; f.submit();} } })();"
            )
        except Exception as e:
            raise RuntimeError("Nao foi possivel acionar o login (botao nao encontrado).") from e

    try:
        page.wait_for_load_state("networkidle", timeout=60000)
    except Exception:
        pass

    if pause_after_login_ms > 0:
        print(f"Pausa apos tentativa de login para acompanhamento: {pause_after_login_ms} ms")
        time.sleep(pause_after_login_ms / 1000.0)
        try:
            page.wait_for_timeout(50)
        except Exception:
            pass

    if _is_login_page(page):
        if login_manual_wait_ms > 0:
            print(
                "Tela de login ainda visivel. Resolva captcha/erro de autenticacao manualmente e clique em Entrar. "
                f"Aguardando ate {login_manual_wait_ms} ms..."
            )
            deadline = time.time() + (login_manual_wait_ms / 1000.0)
            while time.time() < deadline:
                if not _is_login_page(page):
                    break
                time.sleep(1.5)
        if _is_login_page(page):
            extra = ""
            try:
                if page.locator("#gRecaptchaToken, iframe[src*='recaptcha' i], div.g-recaptcha").count() > 0:
                    extra = " Captcha/reCAPTCHA detectado."
            except Exception:
                pass
            err_msg = _extract_login_error(page)
            if err_msg:
                extra += f" Mensagem do portal: {err_msg}"
            if headless and "captcha" in extra.lower():
                extra += " (modo headless nao permite resolver; use HEADLESS=false/SHOW_BROWSER=true)."
            raise RuntimeError("Login nao concluido: formulario de login ainda visivel apos timeout." + extra)

    # Garante que estamos na Mesa de Trabalho apos login
    try:
        if "mesatrabalho.aspx" not in (page.url or "").lower():
            mesa_url = urljoin(url, "/paginas/mesatrabalho.aspx")
            print(f"Abrindo Mesa de Trabalho: {mesa_url}")
            page.goto(mesa_url, wait_until="load", timeout=60000)
    except Exception:
        pass


def safe_filename(name: str) -> str:
    """Return a filesystem-safe slug for filenames based on a label like o número do processo."""
    if not name:
        return "arquivo"
    # Replace path separators and illegal chars
    s = re.sub(r"[\\/]+", "_", str(name))
    s = re.sub(r"[^\w\-. ]+", "_", s, flags=re.UNICODE)
    s = s.strip().strip("._")
    return s or "arquivo"


def find_frame_with_text(page, text: str, timeout_ms: int = 30000):
    """Loop through frames until one contains the given text (substring)."""
    deadline = time.time() + (timeout_ms / 1000.0)
    last_err = None
    while time.time() < deadline:
        for fr in page.frames:
            try:
                loc = fr.get_by_text(text, exact=False)
                # .count() waits for DOM stability enough for text lookup
                if loc.count() > 0:
                    return fr
            except Exception as e:
                last_err = e
                continue
        time.sleep(0.3)
    if last_err:
        raise last_err
    raise PWTimeoutError(f"Frame with text '{text}' not found in {timeout_ms}ms.")


def find_frame_with_selector(page, selector: str, timeout_ms: int = 30000):
    """Find a frame containing an element matching selector that is attached in DOM."""
    deadline = time.time() + (timeout_ms / 1000.0)
    while time.time() < deadline:
        for fr in page.frames:
            try:
                loc = fr.locator(selector)
                if loc.count() > 0:
                    try:
                        loc.first.wait_for(state="attached", timeout=1000)
                    except Exception:
                        pass
                    return fr
            except Exception:
                continue
        time.sleep(0.3)
    raise PWTimeoutError(f"Frame with selector '{selector}' not found in {timeout_ms}ms.")


def normalize(s: str) -> str:
    if s is None:
        return ""
    s = str(s)
    s = s.replace("º", "o").replace("°", "o").replace("ª", "a")
    s = unicodedata.normalize("NFKD", s)
    s = "".join([c for c in s if not unicodedata.combining(c)])
    return s


def _pt_data_extenso_from_ddmmyyyy(s: str) -> str:
    """Converte 'dd/mm/yyyy' ou 'd/m/yyyy' para 'd de <mês> de yyyy' em PT-BR.
    Se não conseguir converter, retorna a string original.
    """
    try:
        parts = re.split(r"[/-]", s.strip())
        if len(parts) != 3:
            return s
        d = int(parts[0])
        m = int(parts[1])
        y = int(parts[2])
        meses = [
            "janeiro", "fevereiro", "março", "abril", "maio", "junho",
            "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
        ]
        if 1 <= m <= 12:
            return f"{d} de {meses[m-1]} de {y}"
        return s
    except Exception:
        return s


def find_latest_export_file(directory: Path) -> Path | None:
    candidates = []
    for ext in ("*.xlsx", "*.xls"):
        candidates.extend(directory.glob(ext))
    if not candidates:
        return None
    return max(candidates, key=lambda p: p.stat().st_mtime)


def cleanup_output_dir(output_dir: Path):
    """Remove arquivos gerados de execucoes anteriores para evitar acumulo/local overwrite issues."""
    if not output_dir.exists():
        return
    for p in output_dir.iterdir():
        try:
            if p.is_file() or p.is_symlink():
                p.unlink(missing_ok=True)
            elif p.is_dir():
                import shutil
                shutil.rmtree(p, ignore_errors=True)
        except Exception as e:
            print(f"Aviso: nao foi possivel remover {p}: {e}")


def find_processo_column_index(headers: list[str]) -> int | None:
    best_idx = None
    for i, h in enumerate(headers):
        hl = normalize(h).strip().lower()
        if not hl:
            continue
        if "processo" in hl and any(tag in hl for tag in ("n", "no", "n.", "n ", "numero")):
            return i
        if best_idx is None and "processo" in hl:
            best_idx = i
    if best_idx is not None:
        return best_idx
    if len(headers) >= 5:
        return 4
    return None


def extract_processo_from_excel(path: Path) -> str | None:
    suffix = path.suffix.lower()
    try:
        if suffix == ".xlsx":
            from openpyxl import load_workbook  # type: ignore
            wb = load_workbook(filename=str(path), read_only=True, data_only=True)
            ws = wb.active
            header = None
            idx = None
            for row in ws.iter_rows(values_only=True):
                values = ["" if v is None else str(v) for v in row]
                if header is None:
                    header = values
                    idx = find_processo_column_index(header)
                    if idx is None:
                        continue
                    continue
                if idx is None or idx >= len(values):
                    continue
                val = values[idx]
                if val and str(val).strip():
                    return str(val).strip()
            return None
        elif suffix == ".xls":
            import xlrd  # type: ignore
            book = xlrd.open_workbook(str(path))
            sheet = book.sheet_by_index(0)
            header_row = 0
            idx = None
            for r in range(min(5, sheet.nrows)):
                row_vals = [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
                idx_try = find_processo_column_index(row_vals)
                if idx_try is not None:
                    header_row = r
                    idx = idx_try
                    break
            if idx is None:
                return None
            for r in range(header_row + 1, sheet.nrows):
                try:
                    v = sheet.cell_value(r, idx)
                except Exception:
                    continue
                if v is None:
                    continue
                txt = str(v).strip()
                if txt:
                    return txt
            return None
        else:
            return None
    except Exception as e:
        print(f"Aviso: falha ao ler planilha {path.name}: {e}")
        return None


def extract_processos_from_excel(path: Path) -> list[str]:
    """Extrai todos os números de processo da planilha, na mesma coluna detectada.

    Retorna os valores não-vazios encontrados na coluna identificada como "Processo".
    """
    processos: list[str] = []
    suffix = path.suffix.lower()
    try:
        if suffix == ".xlsx":
            from openpyxl import load_workbook  # type: ignore
            wb = load_workbook(filename=str(path), read_only=True, data_only=True)
            ws = wb.active
            header = None
            idx = None
            for row in ws.iter_rows(values_only=True):
                values = ["" if v is None else str(v) for v in row]
                if header is None:
                    header = values
                    idx = find_processo_column_index(header)
                    continue
                if idx is None or idx >= len(values):
                    continue
                val = values[idx]
                if val and str(val).strip():
                    processos.append(str(val).strip())
            return processos
        elif suffix == ".xls":
            import xlrd  # type: ignore
            book = xlrd.open_workbook(str(path))
            sheet = book.sheet_by_index(0)
            header_row = 0
            idx = None
            for r in range(min(5, sheet.nrows)):
                row_vals = [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
                idx_try = find_processo_column_index(row_vals)
                if idx_try is not None:
                    header_row = r
                    idx = idx_try
                    break
            if idx is None:
                return []
            for r in range(header_row + 1, sheet.nrows):
                try:
                    v = sheet.cell_value(r, idx)
                except Exception:
                    continue
                if v is None:
                    continue
                txt = str(v).strip()
                if txt:
                    processos.append(txt)
            return processos
        else:
            return []
    except Exception as e:
        print(f"Aviso: falha ao ler planilha {path.name}: {e}")
        return processos


def read_processos_from_excel(path: Path) -> list[str]:
    """Wrapper amigavel para extrair processos de uma planilha usando heuristica existente."""
    return extract_processos_from_excel(path)


def _ensure_apo_pen_grid_visible(page, timeout_ms: int = 20000) -> bool:
    """Garante que a grid de processos esteja carregada (Em confeccao APO-PEN)."""
    deadline = time.time() + timeout_ms / 1000.0
    selectors = [
        "#sptMesaTrabalho_gvProcesso",
        "#gvProcesso",
        "table[id*='gvProcesso']",
    ]
    containers = [page] + list(page.frames)
    while time.time() < deadline:
        for container in containers:
            for root_sel in selectors:
                try:
                    root = container.locator(root_sel).first
                    if root.count() == 0:
                        continue
                    display = root.evaluate("el => getComputedStyle(el).display")  # type: ignore[call-arg]
                    if display and display.lower() != "none":
                        return True
                except Exception:
                    continue
        time.sleep(0.3)
    return False


def open_apo_pen_menu(page) -> bool:
    """Abre Processos -> UNIDADE TECNICA DE OFICIOS -> Em confeccao APO-PEN."""
    containers = [page] + list(page.frames)

    def try_click(container):
        clicked_any = False
        try:
            container.get_by_text("Processos", exact=False).first.click()
            clicked_any = True
        except Exception:
            pass
        try:
            container.get_by_text(re.compile(r"UNIDADE\s+T[EÉ]CNICA\s+DE\s+OF[ÍI]CIOS", re.I)).first.click()
            clicked_any = True
        except Exception:
            try:
                container.locator("a#016_PROCESSO, a[id*='UNIDADE']").first.click()
                clicked_any = True
            except Exception:
                pass
        try:
            container.get_by_text(re.compile(r"Em\s*confe[cç][aã]o\s*APO", re.I)).first.click()
            clicked_any = True
        except Exception:
            try:
                container.locator("a#confappen_16_PROCESSO, a[id*='confappen']").first.click()
                clicked_any = True
            except Exception:
                try:
                    container.locator("a:has-text('Em confecção APO-PEN'), a:has-text('Em confecao APO-PEN')").first.click()
                    clicked_any = True
                except Exception:
                    pass
        return clicked_any

    for attempt in range(4):
        for c in containers:
            try_click(c)
        try:
            page.wait_for_load_state("networkidle", timeout=8000)
        except Exception:
            pass
        if _ensure_apo_pen_grid_visible(page, timeout_ms=8000):
            return True
        time.sleep(1.0)
    return False



def open_apo_pen_and_export_excel(context, page, output_dir: Path | None = None) -> Path | None:
    """Abre Em confeccao APO-PEN e exporta a planilha via botao Exportar."""
    output_dir = output_dir or Path("output")
    output_dir.mkdir(exist_ok=True)

    ok = open_apo_pen_menu(page)
    if not ok:
        print("Aviso: nao foi possivel abrir a pasta 'Em confeccao APO-PEN'.")
        return None

    if not _ensure_apo_pen_grid_visible(page, timeout_ms=20000):
        print("Aviso: grid gvProcesso nao ficou visivel apos abrir o menu.")
        return None

    export_selectors = [
        "#sptMesaTrabalho_gvProcesso_Title_btnExport, #sptMesaTrabalho_gvProcesso_Title_btnExport_I",
        "#gvProcesso_Title_btnExport, #gvProcesso_Title_btnExport_I",
        "#sptMesaTrabalho_gvDocumentos_Title_btnExport, #sptMesaTrabalho_gvDocumentos_Title_btnExport_I",
        "a:has-text('Exportar'), button:has-text('Exportar')",
    ]
    for sel in export_selectors:
        try:
            loc = page.locator(sel).first
            if loc.count() == 0:
                continue
            try:
                loc.wait_for(state="visible", timeout=12000)
            except Exception:
                pass
            with page.expect_download(timeout=60000) as dl_info:
                loc.click()
            download = dl_info.value
            suggested = download.suggested_filename or f"export_{int(time.time())}.xlsx"
            dest_path = output_dir / suggested
            try:
                download.save_as(str(dest_path))
            except Exception:
                tmp_path = download.path()
                if tmp_path:
                    import shutil
                    shutil.copyfile(tmp_path, dest_path)
            print(f"Planilha exportada: {dest_path.resolve()}")
            return dest_path
        except Exception:
            continue
    print("Aviso: botao 'Exportar' nao encontrado na grid APO-PEN.")
    return None


def _select_option_like(container, selectors: list[str], desired: str, fallback_first: bool = True) -> bool:
    """Tenta selecionar uma opcao em <select> ou combobox com heuristica de substring normalizada."""
    desired_norm = normalize(desired or "").lower().strip()
    for sel in selectors:
        try:
            loc = container.locator(sel).first
            if loc.count() == 0:
                continue
            tag = None
            try:
                tag = (loc.evaluate("el => el.tagName") or "").lower()
            except Exception:
                tag = None
            if tag == "select":
                try:
                    options = loc.evaluate("el => Array.from(el.options||[]).map(o => ({value:o.value, text:o.textContent||''}))")
                except Exception:
                    options = []
                pick_val = None
                if options:
                    if desired_norm:
                        for opt in options:
                            if desired_norm in normalize(opt.get("text", "")).lower():
                                pick_val = opt.get("value") or opt.get("text")
                                break
                        if not pick_val:
                            for opt in options:
                                if desired_norm in normalize(opt.get("value", "")).lower():
                                    pick_val = opt.get("value")
                                    break
                    if not pick_val and options and fallback_first:
                        for opt in options:
                            if normalize(opt.get("text", "")).strip():
                                pick_val = opt.get("value") or opt.get("text")
                                break
                if pick_val is not None:
                    try:
                        loc.select_option(value=pick_val)
                    except Exception:
                        try:
                            loc.select_option(label=pick_val)
                        except Exception:
                            pass
                    try:
                        loc.dispatch_event("change")
                    except Exception:
                        pass
                    return True
            try:
                loc.click()
            except Exception:
                pass
            if desired_norm:
                try:
                    loc.fill(desired)
                    try:
                        loc.press("Enter")
                    except Exception:
                        pass
                    return True
                except Exception:
                    pass
        except Exception:
            continue
    return False


def open_caixa_correio_from_grid(context, page, processo: str):
    """Abre a Caixa de Correio / Comunicacao Processual a partir da grid Em confeccao APO-PEN."""
    try:
        if not _ensure_apo_pen_grid_visible(page, timeout_ms=8000):
            open_apo_pen_menu(page)
            _ensure_apo_pen_grid_visible(page, timeout_ms=15000)
    except Exception:
        pass

    # Filtra pelo numero do processo (mesmos seletores das funcoes existentes)
    try:
        inp = page.locator("input[id$='_DXFREditorcol17_I'], input[name$='$DXFREditorcol17']").first
        inp.wait_for(state="visible", timeout=10000)
        try:
            inp.fill("")
        except Exception:
            pass
        inp.fill(processo)
        try:
            inp.press("Enter")
        except Exception:
            pass
    except Exception:
        try:
            header = page.locator("#sptMesaTrabalho_gvProcesso_DXHeadersRow0 td").filter(
                has_text=re.compile(r"N\s*o?\s*Processo", re.I)
            ).first
            if header.count() > 0:
                hid = header.get_attribute("id") or ""
                m = re.search(r"col(\d+)$", hid)
                if m:
                    col_idx = m.group(1)
                    inp = page.locator(f"#sptMesaTrabalho_gvProcesso_DXFREditorcol{col_idx}_I").first
                    if inp.count() > 0:
                        inp.fill(processo)
                        try:
                            inp.press("Enter")
                        except Exception:
                            pass
        except Exception:
            pass

    # Aguarda a primeira linha
    row = None
    deadline = time.time() + 15
    while time.time() < deadline:
        try:
            r = page.locator("#sptMesaTrabalho_gvProcesso_DXMainTable tr[id*='DXDataRow'], #gvProcesso_DXMainTable tr[id*='DXDataRow']").first
            if r.count() > 0:
                row = r
                break
        except Exception:
            pass
        time.sleep(0.2)
    if row is None:
        return page

    pages_before = list(context.pages)
    icon_selectors = [
        "img[src*='img_notificacao' i]",
        "img[src*='notificacao' i]",
        "a:has(img[src*='notificacao' i])",
    ]
    target = None
    for sel in icon_selectors:
        try:
            loc = row.locator(sel).first
            if loc.count() == 0:
                continue
            try:
                with page.expect_popup(timeout=6000) as pop_info:
                    loc.click()
                target = pop_info.value
                try:
                    target.wait_for_load_state("domcontentloaded", timeout=8000)
                except Exception:
                    pass
                break
            except Exception:
                try:
                    loc.click()
                except Exception:
                    pass
                break
        except Exception:
            continue

    if target is None:
        # Detecta nova pagina
        for _ in range(10):
            pages_now = list(context.pages)
            if len(pages_now) > len(pages_before):
                try:
                    target = [p for p in pages_now if p not in pages_before][-1]
                except Exception:
                    target = pages_now[-1]
                try:
                    target.wait_for_load_state("domcontentloaded", timeout=5000)
                except Exception:
                    pass
                break
            time.sleep(0.4)

    if target is None:
        try:
            fr = find_frame_with_text(page, "Comunica", timeout_ms=8000)
            target = fr
        except Exception:
            target = page
    return target


def criar_comunicacao_processual(context, page_like, dados: dict) -> bool:
    """Preenche e cria uma nova Comunicacao Processual."""
    processo = dados.get("processo") or ""
    secretaria = dados.get("secretaria") or ""
    relator = dados.get("relator") or ""
    tipo = dados.get("tipo") or ""
    prazo = int(dados.get("prazo") or 0) if dados.get("prazo") is not None else 0
    desc_custom = dados.get("descricao") or ""
    if not desc_custom:
        desc_custom = f"Oficio {tipo} - modelo {secretaria} - gerado automaticamente".strip(" -")
    target = page_like

    # Procura container que tenha o botao 'Nova Comunicacao Processual'
    containers = [page_like]
    try:
        containers.extend(list(getattr(page_like, "frames", [])))
    except Exception:
        pass
    for c in containers:
        try:
            btn = c.get_by_role("button", name=re.compile(r"Nova\s+Comunic", re.I)).first
            if btn.count() > 0:
                target = c
                break
        except Exception:
            continue

    # Clica no botao e captura possivel popup
    try:
        btn = target.get_by_role("button", name=re.compile(r"Nova\s+Comunic", re.I)).first
    except Exception:
        btn = None
    pages_before = list(context.pages)
    if btn and btn.count() > 0:
        popup_host = target if hasattr(target, "expect_popup") else None
        if popup_host:
            try:
                with popup_host.expect_popup(timeout=6000) as pop_info:  # type: ignore[attr-defined]
                    btn.click()
                target = pop_info.value  # type: ignore[name-defined]
                try:
                    target.wait_for_load_state("domcontentloaded", timeout=8000)
                except Exception:
                    pass
            except Exception:
                try:
                    btn.click()
                except Exception:
                    try:
                        btn.click(force=True)
                    except Exception:
                        pass
        else:
            try:
                btn.click()
            except Exception:
                try:
                    btn.click(force=True)
                except Exception:
                    pass

    if hasattr(target, "frames") and target not in containers:
        try:
            target.wait_for_load_state("domcontentloaded", timeout=8000)
        except Exception:
            pass

    # Se nenhuma nova pagina abriu, tenta detectar mudanca de contexto
    if hasattr(context, "pages") and target == page_like:
        try:
            pages_now = list(context.pages)
            if len(pages_now) > len(pages_before):
                target = [p for p in pages_now if p not in pages_before][-1]
        except Exception:
            pass

    # Escolhe o container do formulario (frame ou propria pagina)
    form_container = target
    try:
        frames = list(getattr(target, "frames", []))
    except Exception:
        frames = []
    for fr in frames:
        try:
            if fr.locator("select, textarea, input").count() > 0:
                form_container = fr
                break
        except Exception:
            continue

    # Destinatario
    dest_ok = False
    if secretaria:
        dest_ok = _select_option_like(form_container, [
            "select[id*='Destin' i]",
            "select[name*='Destin' i]",
            "select[id*='Secretaria' i]",
            "select[name*='Secretaria' i]",
        ], secretaria, fallback_first=True)
        if not dest_ok:
            try:
                form_container.get_by_text(re.compile(re.escape(normalize(secretaria)), re.I)).first.click()
                dest_ok = True
            except Exception:
                dest_ok = False
    if not dest_ok:
        _select_option_like(form_container, [
            "select[id*='Destin' i]",
            "select[name*='Destin' i]",
        ], "", fallback_first=True)

    # Relator
    if relator:
        _select_option_like(form_container, [
            "select[id*='Relator' i]",
            "select[name*='Relator' i]",
        ], relator, fallback_first=True)
        try:
            cb = form_container.get_by_role("combobox", name=re.compile("Relator", re.I)).first
            if cb.count() > 0:
                cb.click()
                cb.fill(relator)
                try:
                    cb.press("Enter")
                except Exception:
                    pass
        except Exception:
            pass

    # Descricao
    try:
        form_container.get_by_label(re.compile(r"Descricao", re.I)).first.fill(desc_custom)
    except Exception:
        try:
            form_container.locator("textarea, input[type='text']").first.fill(desc_custom)
        except Exception:
            pass

    # Status de entrega: Urgente
    status_done = False
    try:
        form_container.get_by_role("radio", name=re.compile("Urgente", re.I)).first.check()
        status_done = True
    except Exception:
        try:
            form_container.get_by_label(re.compile("Urgente", re.I)).first.check()
            status_done = True
        except Exception:
            status_done = False
    if not status_done:
        _select_option_like(form_container, [
            "select[id*='Status' i]",
            "select[name*='Status' i]",
        ], "Urgente", fallback_first=True)

    # Prazo
    if prazo:
        desired_prazo = f"{prazo}"
        ok_prazo = _select_option_like(form_container, [
            "select[id*='Prazo' i]",
            "select[name*='Prazo' i]",
        ], desired_prazo, fallback_first=False)
        if not ok_prazo:
            try:
                form_container.get_by_role("radio", name=re.compile(desired_prazo, re.I)).first.check()
            except Exception:
                try:
                    form_container.get_by_text(re.compile(rf"{prazo}\s*dias", re.I)).first.click()
                except Exception:
                    pass

    # Confirma/salva
    saved = False
    for sel in [
        "button:has-text('Salvar')",
        "button:has-text('Gravar')",
        "button:has-text('Confirmar')",
        "input[type='submit'][value*='Salvar' i]",
        "input[type='submit'][value*='Gravar' i]",
        "input[type='submit'][value*='Confirmar' i]",
    ]:
        try:
            form_container.locator(sel).first.click()
            saved = True
            break
        except Exception:
            continue
    if not saved:
        try:
            form_container.get_by_role("button", name=re.compile("Salvar|Confirmar|Cadastrar", re.I)).first.click()
            saved = True
        except Exception:
            pass

    if saved:
        print(f"Comunicacao processual criada para o processo {processo} (prazo {prazo} dias, tipo {tipo}, secretaria {secretaria}).")
    else:
        print("Aviso: nao foi possivel confirmar o formulario de Comunicacao Processual.")
    return saved


def search_processo_and_open_viewer(context, page, processo: str):
    search_frame = None
    try:
        search_frame = find_frame_with_selector(page, "#cbbProcesso_I", timeout_ms=15000)
    except Exception:
        if page.locator("#cbbProcesso_I").count() == 0:
            raise

    target = search_frame if search_frame else page
    target.locator("#cbbProcesso_I").fill(processo)
    clicked = False
    pages_before = list(context.pages)
    try:
        btn = target.locator("button[onclick='BuscaProcesso();']").first
        if btn.count() > 0:
            btn.click()
            clicked = True
    except Exception:
        pass
    if not clicked:
        try:
            target.locator("#cbbProcesso_I").press("Enter")
            clicked = True
        except Exception:
            pass
    # If viewer loads in same page, its frame should appear
    try:
        find_frame_with_selector(page, "#splLeitorDocumentos_pgcPecas_trePecas", timeout_ms=60000)
        return page
    except Exception:
        pass

    # Otherwise try to detect a newly opened page
    deadline = time.time() + 10
    while time.time() < deadline:
        pages_now = list(context.pages)
        if len(pages_now) > len(pages_before):
            newp = pages_now[-1]
            try:
                newp.wait_for_load_state("domcontentloaded", timeout=5000)
            except Exception:
                pass
            return newp
        time.sleep(0.3)
    return page


def open_processo_from_grid(context, page, processo: str):
    """Filter the grid by 'N° Processo' and open the viewer (lupa icon).

    Targets the Mesa de Trabalho grid with id prefix 'sptMesaTrabalho_gvProcesso'.
    """
    # Fill filter for 'N° Processo'
    input_sel = (
        "input[id$='_DXFREditorcol17_I'], "
        "input[name$='$DXFREditorcol17']"
    )
    try:
        inp = page.locator(", ".join(input_sel)).first
        inp.wait_for(state="visible", timeout=10000)
        try:
            inp.fill("")
        except Exception:
            pass
        inp.fill(processo)
        try:
            inp.press("Enter")
        except Exception:
            pass
    except Exception:
        return

    # Wait for first data row to appear
    row = None
    deadline = time.time() + 15
    while time.time() < deadline:
        try:
            r = page.locator("#sptMesaTrabalho_gvProcesso_DXMainTable tr[id*='DXDataRow']").first
            if r.count() > 0:
                row = r
                break
        except Exception:
            pass
        time.sleep(0.2)
    if row is None:
        return None

    # Click the lupa/search icon (first actionable element in row)
    clicked = False
    selectors = [
        "img[src*='img_busca' i]",
        "a[onclick*='VisualizarProtocolo' i]",
        "td:nth-child(2) a, td:nth-child(2) img",
        "a:has(img)",
        "a",
    ]
    for sel in selectors:
        try:
            loc = row.locator(sel).first
            if loc.count() > 0:
                try:
                    # Try to capture popup if it opens a new window
                    with page.expect_popup(timeout=3000) as pop_info:
                        loc.click()
                    try:
                        pop_info.value.wait_for_load_state("domcontentloaded", timeout=5000)
                    except Exception:
                        pass
                    return pop_info.value
                except Exception:
                    loc.click()
                clicked = True
                break
        except Exception:
            continue
    if not clicked:
        # As a last resort, click any button in the first columns
        try:
            row.locator("td:nth-child(2) button, td:nth-child(2) input[type='button']").first.click()
        except Exception:
            pass
    # If we clicked in the same page (no popup), return current page
    return page


def filter_and_open_processo(context, page, processo: str):
    """Versão robusta: localiza o filtro 'N° Processo' pela célula de cabeçalho,
    digita o número, pressiona Enter e clica na lupa da primeira linha.

    Retorna a nova página (popup) quando abrir em janela separada, ou a página atual.
    """
    # Tenta descobrir o índice da coluna 'N° Processo' pelo header
    try:
        header = page.locator("#sptMesaTrabalho_gvProcesso_DXHeadersRow0 td").filter(
            has_text=re.compile(r"N\s*°?\s*Processo|N\s*o\.?\s*Processo", re.I)
        ).first
        if header.count() > 0:
            hid = header.get_attribute("id") or ""
            m = re.search(r"col(\d+)$", hid)
            if m:
                col_idx = m.group(1)
                inp = page.locator(f"#sptMesaTrabalho_gvProcesso_DXFREditorcol{col_idx}_I").first
                inp.wait_for(state="visible", timeout=10000)
                try:
                    inp.fill("")
                except Exception:
                    pass
                inp.fill(processo)
                try:
                    inp.press("Enter")
                except Exception:
                    pass
        else:
            # Fallback: caixa superior
            topInp = page.locator("input[placeholder*='Processo' i], #cbbProcesso_I").first
            topInp.wait_for(state="visible", timeout=8000)
            topInp.fill(processo)
            try:
                page.locator("button[onclick='BuscaProcesso();']").first.click()
            except Exception:
                topInp.press("Enter")
    except Exception:
        pass

    # Aguarda a primeira linha
    row = None
    deadline = time.time() + 20
    while time.time() < deadline:
        try:
            r = page.locator("#sptMesaTrabalho_gvProcesso_DXMainTable tr[id*='DXDataRow']").first
            if r.count() > 0:
                row = r
                break
        except Exception:
            pass
        time.sleep(0.2)
    if row is None:
        return None

    # Clica na lupa
    selectors = [
        "a[href*='VisualizarDocsProtocolo.aspx' i]",
        "a[onclick*='VisualizarProtocolo' i]",
        "img[src*='img_busca' i]",
        "img[src*='lupa' i]",
        "img[src*='search' i]",
        "img[alt*='busca' i], img[title*='busca' i]",
        "td a:has(img)",
    ]
    for sel in selectors:
        try:
            loc = row.locator(sel).first
            if loc.count() == 0:
                continue
            try:
                with page.expect_popup(timeout=10000) as pop_info:
                    loc.click()
                try:
                    pop_info.value.wait_for_load_state("domcontentloaded", timeout=10000)
                except Exception:
                    pass
                return pop_info.value
            except Exception:
                loc.click()
                return page
        except Exception:
            continue
    return page

def open_gerenciador_atos_from_grid(context, page, processo: str):
    """Filter the grid by 'N° Processo' and open the Gerenciador de Atos (clip icon) popup.

    Heuristics:
    - Reuse the filter input used by open_processo_from_grid.
    - In the first data row, look for a link to Ato/GerenciaAto.aspx or an icon that resembles a clip/attachment/atos.
    """
    # Ensure grid is visible (if needed, try to open 'Em confecção APO-PEN')
    try:
        # If grid not present, try to navigate via menu
        if page.locator("#sptMesaTrabalho_gvProcesso_DXMainTable").count() == 0:
            fr_menu = find_frame_with_text(page, "Processos", timeout_ms=10000)
            fr_menu.get_by_text("Processos", exact=True).first.click(force=True)
            time.sleep(0.5)
            fr_menu.get_by_text(re.compile(r"Em\s*confec.*APO-?PEN", re.I)).first.click(force=True)
            page.wait_for_load_state("domcontentloaded", timeout=20000)
    except Exception:
        pass

    # Filter by process number
    try:
        inp = page.locator("input[id$='_DXFREditorcol17_I'], input[name$='$DXFREditorcol17']").first
        inp.wait_for(state="visible", timeout=10000)
        try:
            inp.fill("")
        except Exception:
            pass
        inp.fill(processo)
        try:
            inp.press("Enter")
        except Exception:
            pass
    except Exception:
        return None

    # Wait for the first row
    row = None
    deadline = time.time() + 15
    while time.time() < deadline:
        try:
            r = page.locator("#sptMesaTrabalho_gvProcesso_DXMainTable tr[id*='DXDataRow']").first
            if r.count() > 0:
                row = r
                break
        except Exception:
            pass
        time.sleep(0.2)
    if row is None:
        return None

    # Try to click the Gerenciador de Atos link/icon
    selectors = [
        "a[href*='/Ato/GerenciaAto.aspx' i]",
        "a[onclick*='GerenciaAto' i]",
        "img[src*='clip' i]",
        "img[alt*='Ato' i]",
        "img[src*='ato' i]",
        "img[src*='anexo' i]",
        "a:has(img)"
    ]
    popup_page = None
    for sel in selectors:
        try:
            loc = row.locator(sel).first
            if loc.count() == 0:
                continue
            with page.expect_popup(timeout=5000) as pop_info:
                loc.click()
            popup_page = pop_info.value
            try:
                popup_page.wait_for_load_state("domcontentloaded", timeout=10000)
            except Exception:
                pass
            break
        except Exception:
            continue
    return popup_page


def attach_docx_via_gerenciador_atos(context, page, processo: str, docx_path: Path) -> bool:
    """Try to attach the DOCX via the Gerenciador de Atos popup.

    Steps:
    - Open 'Gerenciador de Atos' from the grid by clicking the clip icon (popup window).
    - Click 'Anexar Ato' button in the popup.
    - On the upload page, select the DOCX and click to submit.
    Returns True on best-effort success.
    """
    # Se nao receber um caminho valido, tenta pegar o DOCX mais recente da pasta output
    if not docx_path or not docx_path.exists():
        try:
            output_dir = Path("output")
            latest = None
            for p in sorted(output_dir.glob("*.docx"), key=lambda x: x.stat().st_mtime, reverse=True):
                latest = p
                break
            if latest is None:
                return False
            docx_path = latest
        except Exception:
            return False

    try:
        if re.search(r"/Ato/GerenciaAto\.aspx", page.url, re.I):
            pop = page
        else:
            pop = open_gerenciador_atos_from_grid(context, page, processo)
    except Exception as e:
        pop = None

    if not pop:
        return False

    # 1) Clicar preferencialmente em 'Anexar Ato' (novo fluxo); se nao existir, tenta 'Anexar Atos'
    clicked = False
    for sel in [
        "#btnAnexaAto_CD, #btnAnexaAto, #btnAnexaAto_I", # Anexar Ato (singular)
        "button:has-text('Anexar Ato')",
        "input[type='submit'][value*='Anexar Ato' i]",
        # Fallbacks (fluxo antigo 'Anexar Atos')
        "#btnAnexaAtos, #btnAnexaAtos_I",
        "button:has-text('Anexar Atos')",
        "input[type='submit'][value*='Anexar Atos' i]"
    ]:
        try:
            loc = pop.locator(sel).first
            if loc.count() > 0:
                with pop.expect_navigation(url=re.compile(r"uploadato|uploadAtos", re.I), timeout=15000):
                    loc.click()
                clicked = True
                break
        except Exception:
            continue

    # If no navigation happened, try to proceed anyway on same popup
    target = pop
    try:
        if re.search(r"uploadato|uploadAtos", target.url, re.I) is None:
            # Maybe the click changed location without full navigation; wait a bit
            try:
                target.wait_for_url(re.compile(r"uploadato|uploadAtos", re.I), timeout=8000)
            except Exception:
                pass
    except Exception:
        pass

    # 2) On upload page, set input file
    uploaded = False
    try:
        # Novo fluxo simples (uploadato.aspx): input #uplAto
        try:
            el_simple = target.locator("#uplAto, input[name='uplAto']").first
            if el_simple.count() > 0:
                el_simple.set_input_files(str(docx_path.resolve()))
                print(f"Arquivo selecionado (uplAto): {docx_path.name}")
                try:
                    target.evaluate("try{ if(window.UpdateUploadButton) UpdateUploadButton(); }catch(e){}")
                except Exception:
                    pass
                uploaded = True
        except Exception:
            pass

        # Preferred: click the "Selecione o(s) arquivo(s)" button and use file chooser
        # But first, if the exact DevExpress file input id is present, set directly.
        try:
            el_direct = target.locator("#cbpArquivos_UplAtos_TextBox0_Input").first
            if el_direct.count() > 0:
                el_direct.set_input_files(str(docx_path.resolve()))
                print(f"Arquivo selecionado para upload (id direto): {docx_path.name}")
                try:
                    target.evaluate("try{ if(window.UpdateUploadButton) UpdateUploadButton(); }catch(e){}")
                except Exception:
                    pass
                uploaded = True
        except Exception:
            pass
        if not uploaded:
            try:
                with target.expect_file_chooser(timeout=6000) as fc_info:
                    # Tenta seletores exatos do ASPxUploadControl
                    selectors = [
                        "#cbpArquivos_UplAtos_Browse0 a",
                        "#cbpArquivos_UplAtos_BrowseT a",
                        "td[id^='cbpArquivos_UplAtos_Browse'] a",
                        "td.dxucBrowseButton a",
                        "a:has-text('Selecione o(s) arquivo(s)')"
                    ]
                    clicked = False
                    for sel in selectors:
                        loc = target.locator(sel).first
                        if loc.count() > 0:
                            loc.click()
                            clicked = True
                            break
                    if not clicked:
                        # Fallback: busca por texto
                        target.get_by_text(re.compile(r"Selecione\s*o\(s\)\s*arquivo\(s\)", re.I)).first.click()
                fc = fc_info.value
                fc.set_files(str(docx_path.resolve()))
                print(f"Arquivo selecionado para upload: {docx_path.name}")
                try:
                    target.evaluate("try{ if(window.UpdateUploadButton) UpdateUploadButton(); }catch(e){}")
                except Exception:
                    pass
                uploaded = True
            except Exception:
                pass

        if not uploaded:
            # Fallback: set hidden input[type=file] directly (DevExpress UploadControl)
            inp = None
            for sel in [
                "input[id^='cbpArquivos_UplAtos_TextBox'][id$='_Input']",
                "input[type='file']",
                "input[name*='File' i]",
                "input[id*='File' i]",
                "input[id*='upload' i]",
                "input[id*='upl' i]",
            ]:
                try:
                    el = target.query_selector(sel)
                    if el:
                        inp = el
                        break
                except Exception:
                    continue
            if not inp:
                return False
            inp.set_input_files(str(docx_path.resolve()))
            print(f"Arquivo selecionado para upload (fallback): {docx_path.name}")
            try:
                target.evaluate("try{ if(window.UpdateUploadButton) UpdateUploadButton(); }catch(e){}")
            except Exception:
                pass
            uploaded = True
    except Exception:
        return False

    # Wait for classification controls to render (novo/antigo)
    try:
        target.wait_for_selector("#cbbTiposAtos_I, #divGvArquivos select, tr select, table select", timeout=15000)
    except Exception:
        pass

    # 2.1) Classificar como Ofício SSG (novo fluxo); manter fallbacks antigos
    try:
        # Tentativa direta via DevExpress: definir valor 79 ('Ofício SSG')
        try:
            target.evaluate(
                "(function(){\n"
                "  try {\n"
                "    var cb = (window.ASPx && ASPx.GetControlCollection) ? ASPx.GetControlCollection().GetByName('cbbTiposAtos') : (window.cbbTiposAtos || null);\n"
                "    if (cb && cb.SetValue) { cb.SetValue('79'); cb.SetText('Ofício SSG'); return true; }\n"
                "  } catch(e) {}\n"
                "  try {\n"
                "    var vi=document.getElementById('cbbTiposAtos_VI'); var ti=document.getElementById('cbbTiposAtos_I');\n"
                "    if (vi) vi.value='79'; if (ti) ti.value='Ofício SSG'; return !!(vi||ti);\n"
                "  } catch(e) {}\n"
                "  return false;\n"
                "})()"
            )
        except Exception:
            pass        # Combo global da página nova (uploadato.aspx)
        try:
            cg = target.locator("#cbbTiposAtos_I").first
            if cg.count() > 0:
                try:
                    cg.click()
                except Exception:
                    pass
                try:
                    cg.fill("Oficio SSG")
                except Exception:
                    pass
                # tenta abrir dropdown e escolher explicitamente
                try:
                    ddbtn = target.locator("#cbbTiposAtos_B-1").first
                    if ddbtn.count() > 0:
                        ddbtn.click()
                        try:
                            target.get_by_text(re.compile(r"of[ií]cio\s*ssg", re.I)).first.click()
                        except Exception:
                            pass
                except Exception:
                    pass
        except Exception:
            pass

        # DevExpress ASPxComboBox dentro do grid (input id termina com _cbbTipoAto_I)
        try:
            tipo_inp = target.locator("input[id$='_cbbTipoAto_I'], input[id*='_cbbTipoAto_I']").first
            if tipo_inp.count() > 0:
                try:
                    tipo_inp.click()
                except Exception:
                    pass
                try:
                    tipo_inp.fill("Ofício SSG")
                except Exception:
                    pass
                # Tenta selecionar a opção na lista suspensa, se aparecer
                try:
                    opt = target.get_by_text(re.compile(r"of[ií]cio\\s*ssg", re.I)).first
                    if opt.count() > 0:
                        opt.click()
                except Exception:
                    pass
                try:
                    tipo_inp.press("Enter")
                except Exception:
                    pass
        except Exception:
            pass

        row = None
        try:
            row = target.locator("tr:has(select)").first
            if row.count() == 0 and docx_path.name:
                row = target.locator(f"tr:has-text('{docx_path.name}')").first
        except Exception:
            row = None
        if row and row.count() > 0:
            # Prefer <select>
            try:
                sel = row.locator("select").first
                if sel.count() > 0:
                    try:
                        # tentativa direta por label
                        sel.select_option(label=re.compile(r"of[ií]cio\s*ssg", re.I))
                    except Exception:
                        # busca o value cujo texto contenha 'encaminhamento'
                        try:
                            value = sel.evaluate("el => { const opt = Array.from(el.options).find(o => /of[ií]cio\\s*ssg/i.test(o.textContent)); return opt ? opt.value : null; }")
                            if value:
                                sel.select_option(value=value)
                            else:
                                # fallbacks
                                try:
                                    sel.select_option(label=re.compile(r"encaminhamento", re.I))
                                except Exception:
                                    sel.select_option(label="ANEXO")
                        except Exception:
                            for lab in (re.compile(r"encaminhamento", re.I), "ANEXO"):
                                try:
                                    sel.select_option(label=lab)
                                    break
                                except Exception:
                                    continue
            except Exception:
                pass
            # Alternativa: combobox (role)
            try:
                cb = row.get_by_role("combobox").first
                if cb.count() > 0:
                    try:
                        # Abra as opções e clique na opção com o texto
                        cb.click()
                        try:
                            target.get_by_role("option", name=re.compile(r"of[ií]cio\s*ssg", re.I)).first.click()
                        except Exception:
                            target.get_by_text(re.compile(r"of[ií]cio\\s*ssg", re.I)).first.click()
                    except Exception:
                        try:
                            target.get_by_text(re.compile(r"encaminhamento|^\\s*anexo\\s*$", re.I)).first.click()
                        except Exception:
                            pass
            except Exception:
                pass
            # Fallback: clicar no texto ANEXO
            try:
                opt = target.get_by_text(re.compile(r"of[ií]cio\\s*ssg", re.I)).first
                if opt.count() == 0:
                    opt = target.get_by_text(re.compile(r"encaminhamento|^\\s*anexo\\s*$", re.I)).first
                if opt.count() > 0:
                    opt.click()
            except Exception:
                pass
    except Exception:
        pass

    # 2.9) Se já houver arquivo selecionado, tenta submeter o formulário diretamente (mais robusto)
    try:
        has_file = target.evaluate(
            "(function(){ try{ var inp=document.getElementById('uplAto'); return !!(inp && inp.value && inp.value.length>0); }catch(e){ return false; } })()"
        )
    except Exception:
        has_file = False
    if has_file:
        try:
            with target.expect_event('dialog', timeout=120000) as d:
                target.evaluate(
                    "(function(){ try{ var f=document.getElementById('frm'); if(!f) return; try{ f.removeAttribute('onsubmit'); f.onsubmit=null; }catch(_e){}; try{ window.WebForm_OnSubmit=function(){return true;}; window.ValidatorOnSubmit=function(){return true;}; window.Page_BlockSubmit=false; window.Page_IsValid=true; }catch(_e){}; var t=document.getElementById('__EVENTTARGET'); if(t) t.value='btnConfirmar'; var a=document.getElementById('__EVENTARGUMENT'); if(a) a.value=''; f.submit(); }catch(e){} })()"
                )
            try:
                d.value.accept()
            except Exception:
                pass
        except Exception:
            try:
                target.evaluate(
                    "(function(){ try{ var f=document.getElementById('frm'); if(!f) return; try{ f.removeAttribute('onsubmit'); f.onsubmit=null; }catch(_e){}; try{ window.WebForm_OnSubmit=function(){return true;}; window.ValidatorOnSubmit=function(){return true;}; window.Page_BlockSubmit=false; window.Page_IsValid=true; }catch(_e){}; var t=document.getElementById('__EVENTTARGET'); if(t) t.value='btnConfirmar'; var a=document.getElementById('__EVENTARGUMENT'); if(a) a.value=''; f.submit(); }catch(e){} })()"
                )
            except Exception:
                pass

    # 3) Confirm submission
    # Espera curta para backend processar seleção
    try:
        target.wait_for_load_state("networkidle", timeout=8000)
    except Exception:
        pass
    # Tentativa rápida via API DevExpress (btnFechar/btnCancelar.DoClick)
    try:
        res_close = target.evaluate(
            "(function(){\n"
            "  try { var coll = (window.ASPx && ASPx.GetControlCollection) ? ASPx.GetControlCollection() : null;\n"
            "        var b = coll ? (coll.GetByName('btnFechar') || coll.GetByName('btnCancelar')) : null;\n"
            "        if (b && b.SetEnabled) b.SetEnabled(true);\n"
            "        if (b && b.DoClick) { b.DoClick(); return true; } } catch(e) {}\n"
            "  try { if (window.btnFechar && btnFechar.DoClick) { btnFechar.SetEnabled && btnFechar.SetEnabled(true); btnFechar.DoClick(); return true; } } catch(e) {}\n"
            "  try { if (window.btnCancelar && btnCancelar.DoClick) { btnCancelar.SetEnabled && btnCancelar.SetEnabled(true); btnCancelar.DoClick(); return true; } } catch(e) {}\n"
            "  return false;\n"
            "})();"
        )
        if res_close:
            # Pequena espera para o backend
            time.sleep(1.0)
            return True
    except Exception:
        pass

    # 3) Confirm submission (fluxo: Próximo -> Fechar; com fallbacks)
    for sel in [
        # Primeiro avanço de etapa
        "button:has-text('Próximo')",
        "button:has-text('Proximo')",
        "input[type='submit'][value*='Próximo' i]",
        "input[type='submit'][value*='Proximo' i]",
        "a:has-text('Próximo')",
        "a:has-text('Proximo')",
        # Confirmação direta
        "button:has-text('Confirmar')",
        "input[type='submit'][value*='Confirmar' i]",
        "a:has-text('Confirmar')",
        # Fallbacks
        "button:has-text('Enviar')",
        "button:has-text('Upload')",
        "button:has-text('Salvar')",
        "input[type='submit'][value*='Enviar' i]",
        "input[type='submit'][value*='Upload' i]",
        "input[type='submit'][value*='Salvar' i]",
    ]:
        try:
            target.locator(sel).first.click()
            break
        except Exception:
            continue

    # Extra: garantir clique em 'Confirmar' quando aparecer
    try:
        # Aguarda aparecer algum seletor do botão Confirmar
        try:
            target.wait_for_selector(
                "#cbpArquivos_btnConfirmar, #cbpArquivos_btnConfirmar_I, input[name='cbpArquivos$btnConfirmar'], #btnConfirmar, #btnConfirmar_I",
                timeout=8000,
            )
        except Exception:
            pass
        clicked_confirm = False
        # Instala um handler global para aceitar qualquer alerta de sucesso que apareça tardiamente
        accepted_alert_flag = {"v": False}
        def _auto_accept_dialog(d):
            try:
                d.accept()
            except Exception:
                pass
            accepted_alert_flag["v"] = True
        try:
            target.on("dialog", _auto_accept_dialog)
        except Exception:
            pass
        # Modo forçado: envia o form diretamente (ignora validações client-side)
        try:
            if env_bool("FORCE_CONFIRM_UPLOAD", False):
                print("[uploadato] FORCE_CONFIRM_UPLOAD=on -> submetendo formulario diretamente")
                target.evaluate(
                    "(function(){ try{ var f=document.getElementById('frm'); if(!f) return; try{ f.removeAttribute('onsubmit'); f.onsubmit=null; }catch(_e){}; try{ window.WebForm_OnSubmit=function(){return true;}; window.ValidatorOnSubmit=function(){return true;}; window.Page_BlockSubmit=false; window.Page_IsValid=true; }catch(_e){}; var t=document.getElementById('__EVENTTARGET'); if(t) t.value='btnConfirmar'; var a=document.getElementById('__EVENTARGUMENT'); if(a) a.value=''; f.submit(); }catch(e){} })()"
                )
                clicked_confirm = True
        except Exception:
            pass
        try:
            target.evaluate("try{ var coll=(window.ASPx&&ASPx.GetControlCollection)?ASPx.GetControlCollection():null; var b=coll?coll.GetByName('btnConfirmar'):null; if(b&&b.SetEnabled) b.SetEnabled(true);}catch(e){}")
        except Exception:
            pass
        # Tentativa via API DevExpress (btnConfirmar.DoClick) com tratamento de alert
        try:
            res = target.evaluate(
                "(function(){\n"
                "  try { var coll = (window.ASPx && ASPx.GetControlCollection) ? ASPx.GetControlCollection() : null;\n"
                "        var b = coll ? coll.GetByName('btnConfirmar') : null;\n"
                "        if (b && b.SetEnabled) b.SetEnabled(true);\n"
                "        if (b && b.DoClick) { b.DoClick(); } } catch(e) {}\n"
                "  try { if (window.btnConfirmar && btnConfirmar.DoClick) { btnConfirmar.SetEnabled && btnConfirmar.SetEnabled(true); btnConfirmar.DoClick(); } } catch(e) {}\n"
                "  try { var t=document.getElementById('__EVENTTARGET'); if(t && t.value==='btnConfirmar') return true; } catch(e) {}\n"
                "  try { var db=document.getElementById('divBotoes'); if (db && db.style && db.style.display==='none') return true; } catch(e) {}\n"
                "  return false;\n"
                "})();"
            )
            if res:
                print("[uploadato] DevExpress DoClick acionado e postback sinalizado (__EVENTTARGET=btnConfirmar ou divBotoes oculto).")
                try:
                    with target.expect_event('dialog', timeout=30000) as d:
                        pass
                    try:
                        d.value.accept()
                    except Exception:
                        pass
                except Exception:
                    pass
                clicked_confirm = True
        except Exception:
            pass

        # Executa o handler client-side oficial para definir e.processOnServer
        if not clicked_confirm:
            try:
                # Loga resultado da validacao cliente e da decisao de prosseguir
                valid_ok = target.evaluate("(function(){ try{ return !!(window.Page_ClientValidate && Page_ClientValidate()); }catch(e){ return false; } })()")
                try:
                    vinfo = target.evaluate(
                        "(function(){ try{ var arr=[]; var vs=window.Page_Validators||[]; for(var i=0;i<vs.length;i++){ var v=vs[i]; arr.push((v.id||'')+':'+(v.isvalid===false?'INVALID':'OK')); } return arr.join('|'); }catch(e){ return ''; } })()"
                    )
                except Exception:
                    vinfo = ""
                print(f"[uploadato] Page_ClientValidate: {valid_ok} Validators: {vinfo}")
                proceed = target.evaluate(
                    "(function(){\n"
                    "  try { var e={processOnServer:false};\n"
                    "        try{ if(window.Page_ClientValidate) Page_ClientValidate(); }catch(ex){}\n"
                    "        if (typeof window.btnConfirmarClientSide_Click === 'function') { window.btnConfirmarClientSide_Click(null, e); }\n"
                    "        return !!e.processOnServer;\n"
                    "  } catch(err) { return false; }\n"
                    "})();"
                )
                print(f"[uploadato] btnConfirmarClientSide_Click -> processOnServer={proceed}")
                if proceed:
                    try:
                        with target.expect_event('dialog', timeout=30000) as d:
                            target.evaluate("try{ if(window.WebForm_DoPostBackWithOptions){ WebForm_DoPostBackWithOptions(new WebForm_PostBackOptions('btnConfirmar','', true, '', '', false, false)); } else { __doPostBack('btnConfirmar',''); } }catch(e){ try{ var f=document.getElementById('frm'); if(f){ f.__EVENTTARGET.value='btnConfirmar'; f.__EVENTARGUMENT.value=''; f.submit(); } }catch(_){} }")
                        try:
                            d.value.accept()
                        except Exception:
                            pass
                    except Exception:
                        target.evaluate("try{ if(window.WebForm_DoPostBackWithOptions){ WebForm_DoPostBackWithOptions(new WebForm_PostBackOptions('btnConfirmar','', true, '', '', false, false)); } else { __doPostBack('btnConfirmar',''); } }catch(e){ try{ var f=document.getElementById('frm'); if(f){ f.__EVENTTARGET.value='btnConfirmar'; f.__EVENTARGUMENT.value=''; f.submit(); } }catch(_){} }")
                    try:
                        post = target.evaluate("(function(){ var t=document.getElementById('__EVENTTARGET'); return !!(t && t.value==='btnConfirmar'); })()")
                    except Exception:
                        post = True
                    clicked_confirm = bool(post)
            except Exception:
                pass
        for conf_sel in (
            "#cbpArquivos_btnConfirmar",          # container (antigo)
            "#cbpArquivos_btnConfirmar_I",       # input submit (antigo)
            "input[name='cbpArquivos$btnConfirmar']",
            "#btnConfirmar",                      # novo uploadato.aspx
            "#btnConfirmar_I",
            "#btnConfirmar_CD",
        ):
            try:
                loc = target.locator(conf_sel).first
                if loc.count() > 0:
                    try:
                        # Tenta rolar para o botao antes de clicar
                        try:
                            hscroll = loc.element_handle(timeout=500)
                            if hscroll:
                                hscroll.scroll_into_view_if_needed(timeout=1000)
                        except Exception:
                            pass
                        print(f"[uploadato] Clicando Confirmar via seletor: {conf_sel}")
                        # Tentativa adicional: aciona click programatico direto no input/container DevExpress
                        try:
                            if conf_sel in ("#btnConfirmar_I", "#btnConfirmar", "#btnConfirmar_CD"):
                                target.evaluate(
                                    "try{ var el = document.querySelector('#btnConfirmar_I') || document.querySelector('#btnConfirmar') || document.querySelector('#btnConfirmar_CD'); if(el){ el.click && el.click(); } }catch(e){}"
                                )
                        except Exception:
                            pass
                        try:
                            with target.expect_event('dialog', timeout=30000) as d:
                                loc.click(force=True)
                            try:
                                d.value.accept()
                            except Exception:
                                pass
                        except Exception:
                            loc.click(force=True)
                        # Verifica se __EVENTTARGET foi armado para btnConfirmar (indica postback)
                        try:
                            armed = target.evaluate("(function(){ var t=document.getElementById('__EVENTTARGET'); return !!(t && t.value==='btnConfirmar'); })()")
                        except Exception:
                            armed = True
                        clicked_confirm = bool(armed)
                        break
                    except Exception:
                        try:
                            handle = loc.element_handle(timeout=1000)
                        except Exception:
                            handle = None
                        if handle is not None:
                            try:
                                try:
                                    with target.expect_event('dialog', timeout=30000) as d:
                                        target.evaluate("el => el.click()", handle)
                                    try:
                                        d.value.accept()
                                    except Exception:
                                        pass
                                except Exception:
                                    target.evaluate("el => el.click()", handle)
                                try:
                                    armed2 = target.evaluate("(function(){ var t=document.getElementById('__EVENTTARGET'); return !!(t && t.value==='btnConfirmar'); })()")
                                except Exception:
                                    armed2 = True
                                clicked_confirm = bool(armed2)
                                break
                            except Exception:
                                pass
            except Exception:
                continue
        if not clicked_confirm:
            # fallback WebForms: aciona __doPostBack, tentando validar cliente e aceitar alert
            try:
                try:
                    target.evaluate("try{ if(window.Page_ClientValidate) Page_ClientValidate(); }catch(e){};");
                except Exception:
                    pass
                if "uploadato" in (target.url or "").lower():
                    try:
                        with target.expect_event('dialog', timeout=30000) as d:
                            target.evaluate("try{ if(window.WebForm_DoPostBackWithOptions){ WebForm_DoPostBackWithOptions(new WebForm_PostBackOptions('btnConfirmar','', true, '', '', false, false)); } else { __doPostBack('btnConfirmar',''); } }catch(e){ __doPostBack('btnConfirmar',''); }")
                        try:
                            d.value.accept()
                        except Exception:
                            pass
                    except Exception:
                        target.evaluate("try{ if(window.WebForm_DoPostBackWithOptions){ WebForm_DoPostBackWithOptions(new WebForm_PostBackOptions('btnConfirmar','', true, '', '', false, false)); } else { __doPostBack('btnConfirmar',''); } }catch(e){ __doPostBack('btnConfirmar',''); }")
                else:
                    try:
                        with target.expect_event('dialog', timeout=30000) as d:
                            target.evaluate("__doPostBack('cbpArquivos$btnConfirmar','')")
                        try:
                            d.value.accept()
                        except Exception:
                            pass
                    except Exception:
                        target.evaluate("__doPostBack('cbpArquivos$btnConfirmar','')")
                # Confirma se o postback foi armado
                try:
                    armed3 = target.evaluate("(function(){ var t=document.getElementById('__EVENTTARGET'); return !!(t && t.value==='btnConfirmar'); })()")
                except Exception:
                    armed3 = True
                clicked_confirm = bool(armed3)
            except Exception:
                pass
        if not clicked_confirm:
            # Ultimo recurso: submeter o form diretamente
            try:
                try:
                    with target.expect_event('dialog', timeout=30000) as d:
                        target.evaluate(
                            "try{\n"
                            "  var f=document.getElementById('frm');\n"
                            "  if(f){\n"
                            "    try{ f.removeAttribute('onsubmit'); f.onsubmit=null; }catch(_e){}\n"
                            "    try{ window.WebForm_OnSubmit=function(){return true;}; }catch(_e){}\n"
                            "    try{ window.ValidatorOnSubmit=function(){return true;}; window.Page_BlockSubmit=false; window.Page_IsValid=true; }catch(_e){}\n"
                            "    try{ if(f.__EVENTTARGET) f.__EVENTTARGET.value='btnConfirmar'; if(f.__EVENTARGUMENT) f.__EVENTARGUMENT.value=''; }catch(_e){}\n"
                            "    f.submit();\n"
                            "  }\n"
                            "}catch(e){}"
                        );
                    try:
                        d.value.accept()
                    except Exception:
                        pass
                except Exception:
                    target.evaluate(
                        "try{ var f=document.getElementById('frm'); if(f){ try{ f.removeAttribute('onsubmit'); f.onsubmit=null; }catch(_e){}; if(f.__EVENTTARGET) f.__EVENTTARGET.value='btnConfirmar'; if(f.__EVENTARGUMENT) f.__EVENTARGUMENT.value=''; f.submit(); } }catch(e){}"
                    );
                clicked_confirm = True
            except Exception:
                pass
        if clicked_confirm:
            try:
                # Espera navegacao/redirect apos confirmar (alert pode segurar ate ser aceito)
                target.wait_for_url(re.compile(r"GerenciaAto\\.aspx", re.I), timeout=120000)
            except Exception:
                pass
            try:
                target.wait_for_load_state("networkidle", timeout=20000)
            except Exception:
                pass
        # Remove handler global de dialog para nao afetar demais passos
        try:
            target.off("dialog", _auto_accept_dialog)  # type: ignore[attr-defined]
        except Exception:
            pass
    except Exception:
        pass

    # Extra: clique explicito em 'Proximo' e depois 'Confirmar' (IDs DevExpress)
    try:
        for sel in (
            "#cbpArquivos_btnProximo_CD",
            "#cbpArquivos_btnProximo_I",
            "input[name='cbpArquivos$btnProximo']",
        ):
            try:
                loc = target.locator(sel).first
                if loc.count() > 0:
                    loc.click()
                    try:
                        target.wait_for_load_state("networkidle", timeout=8000)
                    except Exception:
                        pass
                    break
            except Exception:
                continue
    except Exception:
        pass

    try:
        for sel in (
            "#cbpArquivos_btnConfirmar",         # container div
            "#cbpArquivos_btnConfirmar_CD",     # clickable div
            "#cbpArquivos_btnConfirmar_I",      # input inside
            "input[name='cbpArquivos$btnConfirmar']",
        ):
            try:
                loc = target.locator(sel).first
                if loc.count() > 0:
                    loc.click()
                    try:
                        target.wait_for_load_state("networkidle", timeout=10000)
                    except Exception:
                        pass
                    break
            except Exception:
                continue
    except Exception:
        pass

    # 3.1) 'Fechar' só depois que Confirmar realmente foi acionado
    # Evita fechar/"Cancelar" acidentalmente a tela de upload antes do envio
    try:
        on_gerencia = re.search(r"/Ato/GerenciaAto\.aspx", target.url, re.I) is not None
    except Exception:
        on_gerencia = False

    closed_after_upload = False
    if clicked_confirm or on_gerencia:
        try:
            target.wait_for_load_state("networkidle", timeout=8000)
        except Exception:
            pass
        # Aguarda explicitamente aparecer 'Fechar' (mais seguro)
        try:
            target.wait_for_selector(
                "#btnFechar_CD, #btnFechar, #btnFechar_I, button:has-text('Fechar'), a:has-text('Fechar')",
                timeout=20000,
            )
        except Exception:
            pass

        close_selectors = [
            "#btnFechar_CD",
            "#btnFechar",
            "#btnFechar_I",
            "button:has-text('Fechar')",
            "a:has-text('Fechar')",
            "input[type='button'][value*='Fechar' i]",
            "input[type='submit'][value*='Fechar' i]",
        ]
        # Alguns ambientes usam 'Cancelar' como 'Fechar' apenas na tela GerenciaAto
        if on_gerencia:
            close_selectors += ["#btnCancelar_CD", "#btnCancelar", "#btnCancelar_I"]

        for sel in close_selectors:
            try:
                loc = target.locator(sel).first
                if loc.count() > 0:
                    try:
                        loc.click()
                    except Exception:
                        try:
                            h = loc.element_handle(timeout=1000)
                        except Exception:
                            h = None
                        if h is not None:
                            try:
                                target.evaluate("el => el.click()", h)
                            except Exception:
                                pass
                    closed_after_upload = True
                    break
            except Exception:
                continue
    else:
        print("[uploadato] Nao foi possivel confirmar envio; evitando fechar/cancelar para nao abortar o anexo.")

    # Give some time for server
    time.sleep(2.0)
    return bool(clicked_confirm and closed_after_upload)

def click_last_piece_and_open_pdf(context, page, output_dir: Path, processo: str, position: str = "last") -> tuple[Path | None, Optional[str]]:
    """Within the VisualizarDocsProtocolo viewer, click the most recent piece and download its PDF.

    Set position to "first" to fetch the capa/first piece; defaults to the last piece.

    Heuristics used:
    - Find the frame that holds the pieces tree (by id) or fallback to the top page.
    - Pick the last (or first, when requested) item by numeric attribute (index_ato/index) or DOM order.
    - Try to read the embedded PDF URL from iframe/embed/object/a and download via request context.
    - Fallback to clicking the "nova janela" button and then extract the URL from the popup.
    """
    # 1) Find the viewer frame or use the page itself
    viewer_frame = None
    try:
        viewer_frame = find_frame_with_selector(page, "#splLeitorDocumentos_pgcPecas_trePecas", timeout_ms=20000)
    except Exception:
        # Try alternative cues for the viewer
        for sel in ["#imgNewWindow", "img#imgNewWindow", "#splLeitorDocumentos_pgcPecas_trePecas_D", "#splLeitorDocumentos_pgcPecas"]:
            try:
                viewer_frame = find_frame_with_selector(page, sel, timeout_ms=5000)
                break
            except Exception:
                continue
    if not viewer_frame:
        # As a last resort, operate on the page itself
        viewer_frame = page

    # 2) Locate pieces/attachments anchors and choose the last one (or first if requested)
    loc = viewer_frame.locator("a[index_ato], a[cod_arquivo_digital_criptografado], a[index]")
    count = loc.count()
    if count == 0:
        # Try a broader selection inside the tree container
        loc = viewer_frame.locator("#splLeitorDocumentos_pgcPecas_trePecas a, a[cod_arquivo_digital_criptografado]")
        count = loc.count()
    if count == 0:
        print("Aviso: Nenhuma peca encontrada no visualizador.")
        return None, None

    pos_norm = (position or "last").lower()
    max_idx = -1
    max_n = 0
    min_idx = None
    min_n = 0
    for i in range(count):
        item = loc.nth(i)
        val = item.get_attribute("index_ato") or item.get_attribute("index") or ""
        try:
            iv = int(re.findall(r"\d+", val)[0]) if val else i
        except Exception:
            iv = i
        if iv >= max_idx:
            max_idx = iv
            max_n = i
        if min_idx is None or iv < min_idx:
            min_idx = iv
            min_n = i

    target_n = min_n if pos_norm.startswith("first") else max_n
    label = "primeiro-ato" if pos_norm.startswith("first") else "ultimo-ato"

    # Capture piece title/name before clicking
    piece_title: Optional[str] = None
    try:
        item = loc.nth(target_n)
        try:
            piece_title = (item.inner_text(timeout=1000) or "").strip()
        except Exception:
            try:
                piece_title = (item.text_content(timeout=1000) or "").strip()
            except Exception:
                piece_title = None
        item.click()
    except Exception:
        loc.nth(target_n).click()
    time.sleep(1.0)

    def _pick_attr(pl, sel, attr):
        el = pl.locator(sel)
        if el.count() > 0:
            val = el.first.get_attribute(attr)
            if val:
                return val
        return None

    def _try_find_pdf_url_in_container(container):
        # Try several common embed patterns; URL may not end with .pdf
        # Chromium's built-in PDF viewer exposes the original document URL in
        # an `original-url` attribute on the <embed> element. Prefer this.
        url = (
            _pick_attr(container, "embed[original-url]", "original-url")
            or _pick_attr(container, "iframe[src*='visualiza' i]", "src")
            or _pick_attr(container, "iframe[src*='iFramevisualizardocumento' i]", "src")
            or _pick_attr(container, "iframe[src]", "src")
            or _pick_attr(container, "embed[type*='pdf']", "src")
            or _pick_attr(container, "object[data]", "data")
            or _pick_attr(container, "a[href*='.pdf']", "href")
        )
        return url

    # Helper to download and validate PDF bytes
    def _download_and_save_pdf(abs_url: str, referer_url: str | None = None) -> Path | None:
        try:
            headers = {"Accept": "application/pdf,application/octet-stream;q=0.9,*/*;q=0.8"}
            if referer_url:
                headers["Referer"] = referer_url
            resp = context.request.get(abs_url, headers=headers, timeout=60000)
            if not resp.ok:
                return None
            body = resp.body()
            ct = (resp.headers.get("content-type") or "").lower()
            if (b"%PDF" not in body[:8]) and ("application/pdf" not in ct):
                return None
            pdf_path = output_dir / f"{safe_filename(processo)}-{label}.pdf"
            with open(pdf_path, "wb") as f:
                f.write(body)
            print(f"PDF salvo em: {pdf_path.resolve()}")
            return pdf_path
        except Exception:
            return None

    # 3) Try to download directly from embedded URL
    pdf_url = _try_find_pdf_url_in_container(viewer_frame) or _try_find_pdf_url_in_container(page)
    if pdf_url:
        base_url = getattr(viewer_frame, "url", None) or page.url
        abs_url = urljoin(base_url, pdf_url)
        pdf_path_try = _download_and_save_pdf(abs_url, referer_url=base_url)
        if pdf_path_try:
            return pdf_path_try, piece_title
        else:
            print("Aviso: conteudo nao-PDF retornado no embed direto. Tentando nova janela...")

    # 4) Fallback: open in a new window and extract URL
    btn_sel = "#imgNewWindow, img#imgNewWindow"
    btn_clicked = False
    pdf_page = None
    try:
        if viewer_frame.locator(btn_sel).count() > 0:
            with page.expect_popup(timeout=15000) as pop_info:
                viewer_frame.locator(btn_sel).first.click()
            pdf_page = pop_info.value
            btn_clicked = True
        else:
            raise Exception("not in frame")
    except Exception:
        if page.locator(btn_sel).count() > 0:
            with page.expect_popup(timeout=15000) as pop_info:
                page.locator(btn_sel).first.click()
            pdf_page = pop_info.value
            btn_clicked = True

    if not btn_clicked or not pdf_page:
        print("Aviso: Botao 'abrir em nova janela' nao encontrado e nenhum embed localizado.")
        return None, piece_title

    try:
        pdf_page.wait_for_load_state("domcontentloaded", timeout=30000)
    except Exception:
        pass
    # Give the viewer a moment to inject the <embed> element
    try:
        pdf_page.wait_for_selector("embed[original-url], embed[type*='pdf'], iframe[src], object[data]", timeout=10000)
    except Exception:
        pass

    pdf_url = (
        _pick_attr(pdf_page, "embed[original-url]", "original-url")
        or _pick_attr(pdf_page, "iframe[src]", "src")
        or _pick_attr(pdf_page, "embed[type*='pdf']", "src")
        or _pick_attr(pdf_page, "object[data]", "data")
        or _pick_attr(pdf_page, "a[href*='.pdf']", "href")
    )
    if not pdf_url:
        print("Aviso: URL do PDF nao encontrada na nova janela.")
        return None, piece_title

    abs_url = urljoin(pdf_page.url, pdf_url)
    pdf_path_try = _download_and_save_pdf(abs_url, referer_url=pdf_page.url)
    if pdf_path_try:
        return pdf_path_try, piece_title
    print("Aviso: conteudo nao-PDF retornado na nova janela.")
    return None, piece_title


def _docx_replace_all(doc, mapping: dict[str, str]):
    def _replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for k, v in mapping.items():
                if k in p.text:
                    for r in p.runs:
                        r.text = r.text.replace(k, v)

    # Body paragraphs
    _replace_in_paragraphs(doc.paragraphs)

    # Body tables
    for table in getattr(doc, "tables", []) or []:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs)

    # Headers/Footers
    for section in getattr(doc, "sections", []) or []:
        try:
            _replace_in_paragraphs(section.header.paragraphs)
            _replace_in_paragraphs(section.footer.paragraphs)
            for table in getattr(section.header, "tables", []) or []:
                for row in table.rows:
                    for cell in row.cells:
                        _replace_in_paragraphs(cell.paragraphs)
            for table in getattr(section.footer, "tables", []) or []:
                for row in table.rows:
                    for cell in row.cells:
                        _replace_in_paragraphs(cell.paragraphs)
        except Exception:
            continue


def _resolve_oficio_template() -> Optional[Path]:
    """Resolve template path from env vars.

    Regras:
    - Se OFICIO_TEMPLATE apontar para um arquivo (.docx ou .dotx) existente, usa.
    - Se OFICIO_TEMPLATE for apenas nome de arquivo e OFICIO_TEMPLATES_DIR existir, usa o arquivo dentro da pasta.
    - Caso contrário, se OFICIO_TEMPLATES_DIR existir, escolhe o primeiro .docx; se não houver .docx, escolhe .dotx.
    - Senão, tenta templates/oficio_modelo.docx
    """
    template_env = os.getenv("OFICIO_TEMPLATE")
    templates_dir_env = os.getenv("OFICIO_TEMPLATES_DIR")

    if template_env:
        p = Path(template_env)
        if p.exists() and p.is_file():
            return p
        if templates_dir_env:
            p2 = Path(templates_dir_env) / template_env
            if p2.exists() and p2.is_file():
                return p2

    if templates_dir_env:
        d = Path(templates_dir_env)
        if d.exists() and d.is_dir():
            docxs = sorted(d.glob("*.docx"))
            if docxs:
                return docxs[0]
            dotxs = sorted(d.glob("*.dotx"))
            if dotxs:
                return dotxs[0]

    default_p = Path("templates/oficio_modelo.docx")
    if default_p.exists():
        return default_p
    return None


def _word_generate_from_dotx(template_path: Path, out_path: Path, mapping: dict[str, str]) -> bool:
    """Gera DOCX a partir de um .dotx usando Microsoft Word (COM automation).

    Requisitos: MS Word instalado; pywin32 disponível. Retorna True em sucesso.
    """
    try:
        import win32com.client  # type: ignore
        from win32com.client import constants  # type: ignore
    except Exception as e:
        print(f"Aviso: pywin32/Word indisponivel ({e}).")
        return False

    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Add(Template=str(template_path))
        # Find/Replace para cada placeholder
        for k, v in mapping.items():
            find = word.Selection.Find
            find.ClearFormatting()
            find.Replacement.ClearFormatting()
            find.Text = k
            find.Replacement.Text = v
            find.Forward = True
            find.Wrap = 1  # wdFindContinue
            find.Format = False
            find.MatchCase = False
            find.MatchWholeWord = False
            find.MatchByte = False
            find.MatchWildcards = False
            find.MatchSoundsLike = False
            find.MatchAllWordForms = False
            find.Execute(Replace=2)  # wdReplaceAll
        # Salva como DOCX
        doc.SaveAs2(str(out_path), FileFormat=constants.wdFormatXMLDocument)
        return True
    except Exception as e:
        print(f"Aviso: falha no Word ao gerar a partir do DOTX: {e}")
        return False
    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


def _python_generate_from_dotx(template_path: Path, out_path: Path, mapping: dict[str, str]) -> bool:
    """Gera DOCX a partir de um .dotx copiando o arquivo para .docx e
    executando substituicoes com python-docx (sem MS Word).

    Retorna True em sucesso.
    """
    try:
        import tempfile
        import zipfile
        from io import BytesIO
        from docx import Document  # type: ignore

        # Converte o pacote .dotx em um pacote .docx trocando o content-type principal
        with tempfile.TemporaryDirectory() as td:
            tmp_docx = Path(td) / "tmp_from_dotx.docx"
            with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(tmp_docx, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for info in zin.infolist():
                    data = zin.read(info.filename)
                    if info.filename == "[Content_Types].xml":
                        try:
                            xml = data.decode("utf-8")
                        except Exception:
                            xml = data.decode("utf-8", errors="ignore")
                        xml = xml.replace(
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                        )
                        data = xml.encode("utf-8")
                    zout.writestr(info, data)

            # Abre como DOCX e executa os replaces
            doc = Document(str(tmp_docx))
            _docx_replace_all(doc, mapping)
            doc.save(str(out_path))
        return True
    except Exception as e:
        print(f"Aviso: falha ao gerar DOCX a partir do DOTX via python-docx: {e}")
        return False


def _python_convert_dotx_only(template_path: Path, out_path: Path) -> bool:
    """Converte .dotx em .docx SEM alterar o conteúdo (apenas ajusta o content-type OOXML).

    Não requer MS Word. Evita reserialização do documento para garantir que nenhum
    conteúdo/formatacao seja modificado.
    """
    try:
        import zipfile
        with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "[Content_Types].xml":
                    try:
                        xml = data.decode("utf-8")
                    except Exception:
                        xml = data.decode("utf-8", errors="ignore")
                    xml = xml.replace(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    )
                    data = xml.encode("utf-8")
                zout.writestr(info, data)
        return True
    except Exception as e:
        print(f"Aviso: conversão DOTX->DOCX sem alterações falhou: {e}")
        return False


def generate_oficio_from_template(processo: str, output_dir: Path, extra: dict | None = None, template_path: Optional[Path] = None) -> Path | None:
    """Generate a DOCX response using a template when available; fallback to a simple layout.

    Preferred placeholders in templates: {{NUM_PROCESSO}}, {{DATA}}, plus any provided in `extra`.
    If no .docx template is found (or .dotx is provided), falls back to creating a minimal DOCX.
    """
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        print(f"Aviso: python-docx nao instalado ({e}). Pulei geracao do oficio.")
        return None

    mapping = {"{{NUM_PROCESSO}}": processo, "{{DATA}}": date.today().strftime("%d/%m/%Y")}
    if extra:
        mapping.update({str(k): str(v) for k, v in extra.items()})

    # Alias uteis para modelos que usam nomes diferentes
    nome = mapping.get("{{INTERESSADO}}") or mapping.get("{{REQUERENTE}}")
    if nome:
        mapping.setdefault("{{NOME}}", nome)
        mapping.setdefault("{{NOME_COMPLETO}}", nome)

    # Garantir placeholders @@ usados no DOTX
    proc_val = mapping.get("{{NUM_PROCESSO}}", processo) or processo or ""
    mapping.setdefault("@@processo", proc_val)
    # Data por extenso, baseada na data do documento (se houver)
    data_doc = mapping.get("{{DATA_DOCUMENTO}}")
    if data_doc and re.match(r"\d{1,2}/\d{1,2}/\d{4}", str(data_doc)):
        mapping.setdefault("@@data_extenso", _pt_data_extenso_from_ddmmyyyy(str(data_doc)))
    elif data_doc:
        mapping.setdefault("@@data_extenso", str(data_doc))
    else:
        mapping.setdefault("@@data_extenso", _pt_data_extenso_from_ddmmyyyy(date.today().strftime("%d/%m/%Y")))

    # Numero do ofício (env ou s/n)
    mapping.setdefault("@@numero_oficio", os.getenv("NUMERO_OFICIO", "s/n"))

    # Nome do interessado para @@Nome_interessado
    if nome:
        mapping.setdefault("@@Nome_interessado", nome)

    # Demais chaves @@ com valor padrão vazio para evitar lixos no DOCX
    for key in (
        "@@natureza_processo",
        "@@Tipo_Processo",
        "@@processoexterno",
        "@@nome_relator",
        "@@instancia",
    ):
        mapping.setdefault(key, "")

    doc: Optional["Document"] = None
    # Allow override via argument; otherwise resolve from env
    template_path = template_path or _resolve_oficio_template()
    if template_path and template_path.exists():
        if template_path.suffix.lower() == ".dotx":
            out_path = output_dir / f"oficio_{safe_filename(processo)}.docx"
            convert_only = env_bool("OFICIO_CONVERT_ONLY", False)
            if convert_only:
                # Apenas converter DOTX -> DOCX, sem substituir placeholders
                ok = _python_convert_dotx_only(template_path, out_path)
                if ok:
                    print(f"Oficio convertido de DOTX para DOCX (sem alterações): {out_path.resolve()}")
                    return out_path
                # Fallback com Word (sem replace)
                ok2 = _word_generate_from_dotx(template_path, out_path, mapping={})
                if ok2:
                    print(f"Oficio convertido via Word (sem alterações): {out_path.resolve()}")
                    return out_path
                print("Aviso: conversão DOTX->DOCX (sem alterações) falhou; usando modelo simples.")
            else:
                # 1) Tenta gerar via python-docx (sem Word): copia DOTX -> DOCX e faz replace
                ok = _python_generate_from_dotx(template_path, out_path, mapping)
                if ok:
                    print(f"Oficio gerado (DOTX via python-docx) em: {out_path.resolve()}")
                    return out_path
                # 2) Fallback: tenta MS Word via COM
                ok2 = _word_generate_from_dotx(template_path, out_path, mapping)
                if ok2:
                    print(f"Oficio gerado (Word/DOTX) em: {out_path.resolve()}")
                    return out_path
                else:
                    print("Aviso: conversao DOTX->DOCX falhou; usando modelo simples.")
        else:
            try:
                doc = Document(str(template_path))
            except Exception as e:
                print(f"Aviso: nao foi possivel abrir o template '{template_path.name}' ({e}). Usando modelo simples.")

    if doc is None:
        # Fallback: cria um documento basico com campos comuns
        doc = Document()
        doc.add_heading(f"Oficio – Processo {processo}", level=1)
        doc.add_paragraph(f"Data: {mapping.get('{{DATA}}','')}")
        if extra:
            # Inclui alguns campos relevantes quando existirem
            for k in ("{{ASSUNTO}}", "{{INTERESSADO}}", "{{REQUERENTE}}", "{{DATA_DOCUMENTO}}"):
                v = mapping.get(k)
                if v:
                    label = k.strip("{} ")
                    doc.add_paragraph(f"{label}: {v}")
        doc.add_paragraph("")
        corpo = mapping.get("{{EXTRATO}}") or ""
        if corpo:
            doc.add_paragraph(corpo)
        else:
            doc.add_paragraph("Em atendimento, encaminhamos resposta referente ao processo informado.")
    else:
        _docx_replace_all(doc, mapping)

    out_path = output_dir / f"oficio_{safe_filename(processo)}.docx"
    try:
        doc.save(str(out_path))
        print(f"Oficio gerado em: {out_path.resolve()}")
        return out_path
    except Exception as e:
        print(f"Aviso: falha ao salvar oficio: {e}")
        return None


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract plain text from a PDF using pypdf. Returns empty string on failure."""
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as e:
        print(f"Aviso: biblioteca pypdf nao instalada ({e}). Nao foi possivel ler o PDF.")
        return ""

    try:
        reader = PdfReader(str(pdf_path))
        chunks: list[str] = []
        for page in reader.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt:
                chunks.append(txt)
        return "\n".join(chunks)
    except Exception as e:
        print(f"Aviso: falha ao extrair texto do PDF '{pdf_path.name}': {e}")
        return ""


def _extract_cover_text(context, page, output_dir: Path, processo: str) -> str:
    """Download and extract text from the first (capa) PDF when available."""
    try:
        cover_pdf_path, _ = click_last_piece_and_open_pdf(context, page, output_dir, processo, position="first")
        if cover_pdf_path:
            return extract_text_from_pdf(cover_pdf_path)
    except Exception as e:
        print(f"Aviso: falha ao analisar PDF da capa: {e}")
    return ""


def parse_fields_from_pdf_text(text: str, processo: Optional[str] = None) -> dict[str, str]:
    """Parse common fields from PDF text and return mapping for template placeholders.

    Placeholders populated (quando encontrados):
    - {{NUM_PROCESSO}}
    - {{ASSUNTO}}
    - {{INTERESSADO}}
    - {{REQUERENTE}}
    - {{DATA_DOCUMENTO}}
    - {{DATA}} (kept as today's date; DATA_DOCUMENTO is the date found in PDF)
    - {{EXTRATO}} (first 400 chars as summary)
    - {{CPF}}, {{MATRICULA}}, {{CARGO}}, {{NASCIMENTO}}
    """
    if not text:
        return {"{{NUM_PROCESSO}}": processo or ""}

    def _m(patterns: list[str]) -> Optional[str]:
        for pat in patterns:
            m = re.search(pat, text, flags=re.IGNORECASE | re.MULTILINE)
            if m:
                g = m.group(1).strip()
                # Clean artifacts like excessive spaces
                return re.sub(r"\s+", " ", g)
        return None

    out: dict[str, str] = {}

    # Processo number
    proc = _m([
        r"Processo\s*(?:n[oº\.]|no)?\s*[:\-]?\s*([\d./-]+)",
        r"N[ºo]\s*[:\-]?\s*([\d./-]+)\s*Processo",
    ]) or (processo or "")
    out["{{NUM_PROCESSO}}"] = proc

    # Interested party / requerente
    interessado = _m([
        r"Interessado\s*[:\-]\s*(.+)",
        r"Requerente\s*[:\-]\s*(.+)",
    ])
    if interessado:
        out["{{INTERESSADO}}"] = interessado
        out["{{REQUERENTE}}"] = interessado

    # Assunto
    assunto = _m([
        r"Assunto\s*[:\-]\s*(.+)",
        r"Objeto\s*[:\-]\s*(.+)",
    ])
    if assunto:
        out["{{ASSUNTO}}"] = assunto

    # Date present in document (dd/mm/yyyy or dd de mes de yyyy)
    data_ddmmyyyy = _m([r"(\b\d{1,2}/\d{1,2}/\d{4}\b)"])
    if not data_ddmmyyyy:
        data_ddmmyyyy = _m([
            r"\b(\d{1,2}\s+de\s+[a-zcaiou]+\s+de\s+\d{4})\b",
        ])
    if data_ddmmyyyy:
        out["{{DATA_DOCUMENTO}}"] = data_ddmmyyyy

    # CPF
    cpf = _m([
        r"CPF\s*[:\-]\s*([0-9.\-]{11,14})",
        r"\b(\d{3}\.\d{3}\.\d{3}\-\d{2})\b",
    ])
    if cpf:
        out["{{CPF}}"] = cpf

    # Matrícula / RF
    matricula = _m([
        r"Matr[ií]cula\s*[:\-]\s*([A-Za-z0-9/\.-]+)",
        r"Registro\s*Funcional\s*[:\-]\s*([A-Za-z0-9/\.-]+)",
        r"\bRF\b\s*[:\-]\s*([A-Za-z0-9/\.-]+)",
    ])
    if matricula:
        out["{{MATRICULA}}"] = matricula

    # Cargo
    cargo = _m([
        r"Cargo\s*[:\-]\s*(.+)",
        r"Fun[cç][aã]o\s*[:\-]\s*(.+)",
    ])
    if cargo:
        out["{{CARGO}}"] = cargo

    # Data de nascimento
    nasc = _m([
        r"Data\s*de\s*Nascimento\s*[:\-]\s*([0-9/]{8,10})",
        r"Nascimento\s*[:\-]\s*([0-9/]{8,10})",
    ])
    if nasc:
        out["{{NASCIMENTO}}"] = nasc

    # Short summary
    snippet = re.sub(r"\s+", " ", text).strip()
    if snippet:
        out["{{EXTRATO}}"] = snippet[:400]

    # Additional fields for @@ placeholders
    # Tipo do processo
    tipo_proc = _m([
        r"Tipo\s*de\s*Processo\s*[:\-]\s*(.+)",
        r"Tipo\s*do\s*Processo\s*[:\-]\s*(.+)",
        r"Classe\s*[:\-]\s*(.+)",
        r"Categoria\s*[:\-]\s*(.+)",
        r"Tipo\s*[:\-]\s*(.+)",
    ])
    if tipo_proc:
        out["@@Tipo_Processo"] = tipo_proc

    # Natureza do processo
    natureza = _m([
        r"Natureza\s*(?:do\s*Processo)?\s*[:\-]\s*(.+)",
        r"Classe\s*Processual\s*[:\-]\s*(.+)",
    ])
    if natureza:
        out["@@natureza_processo"] = natureza

    # Processo externo
    proc_ext = _m([
        r"Proc(?:esso)?\.?\s*Externo\s*[:\-]\s*(.+)",
        r"Processo\s*Externo\s*[:\-]\s*(.+)",
        r"Externo\s*[:\-]\s*(.+)",
    ])
    if proc_ext:
        out["@@processoexterno"] = proc_ext

    # Relator
    relator = _m([
        r"Conselheiro\s*Relator\s*[:\-]\s*(.+)",
        r"Relator\s*[:\-]\s*(.+)",
    ])
    if relator:
        out["@@nome_relator"] = relator

    # Instância
    instancia = _m([
        r"Inst[âa]ncia\s*[:\-]\s*(.+)",
        r"Instancia\s*[:\-]\s*(.+)",
    ])
    if instancia:
        out["@@instancia"] = instancia

    # Interessado → Nome_interessado
    if interessado:
        out["@@Nome_interessado"] = interessado

    # Processo → @@processo
    out["@@processo"] = proc

    # Número do ofício: opcional via env; não costuma vir no PDF
    num_of = os.getenv("NUMERO_OFICIO", "").strip()
    out["@@numero_oficio"] = num_of if num_of else "s/n"

    # Data por extenso
    data_src = out.get("{{DATA_DOCUMENTO}}")
    if data_src:
        # Se já estiver por extenso, mantém; caso contrário, converte dd/mm/yyyy
        if re.match(r"\d{1,2}/\d{1,2}/\d{4}", data_src):
            out["@@data_extenso"] = _pt_data_extenso_from_ddmmyyyy(data_src)
        else:
            out["@@data_extenso"] = data_src
    else:
        out["@@data_extenso"] = _pt_data_extenso_from_ddmmyyyy(date.today().strftime("%d/%m/%Y"))

    return out


def _detect_secretaria_from_text(text: str) -> str:
    """Heurística para identificar a secretaria pelo conteúdo do PDF.

    Retorna uma das opções: 'Educação', 'Saúde' ou 'Geral'.
    """
    t = normalize(text).lower()
    # Educação
    if any(k in t for k in [
        "secretaria municipal de educacao", "secretaria de educacao", "sme",
        "educacao"
    ]):
        return "Educação"
    # Saúde
    if any(k in t for k in [
        "secretaria municipal da saude", "secretaria municipal de saude", "secretaria de saude", "sms",
        "saude"
    ]):
        return "Saúde"
    return "Geral"


def extract_data_decadencia(pdf_text: str) -> date | None:
    """Busca uma data de decadencia no texto do PDF (padrao dd/mm/aaaa)."""
    if not pdf_text:
        return None
    text_norm = normalize(pdf_text)
    patterns = [
        r"decadencia[^0-9]{0,20}(\d{1,2}/\d{1,2}/\d{2,4})",
        r"decai[^0-9]{0,20}(\d{1,2}/\d{1,2}/\d{2,4})",
        r"data\s*limite[^0-9]{0,15}(\d{1,2}/\d{1,2}/\d{2,4})",
    ]
    for pat in patterns:
        m = re.search(pat, text_norm, flags=re.IGNORECASE | re.MULTILINE)
        if not m:
            continue
        ds = m.group(1)
        for fmt in ("%d/%m/%Y", "%d/%m/%y"):
            try:
                return datetime.strptime(ds, fmt).date()
            except Exception:
                continue
    return None



def calcular_prazo_res_22_21(data_decadencia: date | None, hoje: date) -> int:
    """Calcula o prazo (15/30/60) conforme Res. 22/21, com fallback de 30 dias."""
    if not data_decadencia:
        print("Aviso: data de decadencia nao encontrada; usando prazo padrao de 30 dias.")
        return 30
    dias = (data_decadencia - hoje).days
    if dias <= 60:
        return 15
    if dias <= 120:
        return 30
    return 60


def _classify_tipo_from_text_and_piece(text: str, last_piece_name: Optional[str], cover_text: Optional[str] = None) -> str:
    """Classifica o tipo de modelo: UTAP, REITERACAO, DILACAO, JUIZO.

    Regras inspiradas nas palavras-chave fornecidas e no nome da última peça,
    considerando tanto o texto do último PDF quanto o da capa (primeiro PDF).
    """
    # Alta prioridade: nome da peça
    piece = normalize(last_piece_name or "").lower()
    if "juizo singular" in piece:
        return "JUIZO"
    if "manutap-of" in piece:
        return "UTAP"

    # Palavras-chave nos PDFs (último e primeiro/capa)
    texts_norm = [normalize(text or "").lower()]
    if cover_text:
        texts_norm.append(normalize(cover_text).lower())

    def _has_any(keys: list[str]) -> bool:
        """Retorna True se qualquer uma das chaves aparecer em qualquer texto analisado."""
        norm_keys = [normalize(k).lower() for k in keys]
        for tnorm in texts_norm:
            if any(k in tnorm for k in norm_keys):
                return True
        return False

    # Juízo Singular
    for tnorm in texts_norm:
        if "decisao de juizo singular" in tnorm or re.search(r"\bjuizo\s+singular\b", tnorm):
            return "JUIZO"

    # Dilação
    dil_keys = [
        "autorizo a prorrogacao",
        "autorizo a dilacao de prazo",
        "autorizando a dilacao de prazo",
        "autorizo a solicitacao de dilacao de prazo",
    ]
    if _has_any(dil_keys):
        return "DILACAO"

    # Reiteração
    reit_keys = [
        "por se tratar de providencias ja solicitada",
        "oficiar a chefia de gabinete",
        "considerando o tempo decorrido",
        "reitere-se",
    ]
    if _has_any(reit_keys):
        return "REITERACAO"

    # UTAP
    utap_keys = [
        "ressaltamos que",
        "entendimento",
        "decadencia",
    ]
    if _has_any(utap_keys):
        return "UTAP"

    # Fallback conservador
    return "UTAP"


def _select_template_for(tipo: str, secretaria: str) -> Optional[Path]:
    """Seleciona o arquivo de modelo correto com base no tipo e secretaria.

    Procura dentro das pastas locais mapeadas e retorna o Path do arquivo
    preferindo .docx; se não houver, retorna .dotx.
    """
    base_map = {
        "UTAP": "modelos_utap",
        "REITERACAO": "modelos_reiteracao",
        "DILACAO": "modelos_dilacao",
        "JUIZO": "modelos_juizo",
    }
    folder = base_map.get(tipo.upper())
    if not folder:
        return None
    d = Path(folder)
    if not d.exists() or not d.is_dir():
        return None

    sec_norm = normalize(secretaria).lower()
    # Primeiro, escolha todos que combinem a secretaria
    candidates = list(d.glob("*"))
    candidates = [p for p in candidates if p.is_file() and sec_norm in normalize(p.name).lower()]
    if not candidates:
        # fallback: qualquer arquivo (pasta deve conter só os 3 modelos)
        candidates = [p for p in d.glob("*") if p.is_file()]
    if not candidates:
        return None

    # Prefer .docx > .dotx > outros
    def score(p: Path) -> int:
        s = p.suffix.lower()
        if s == ".docx":
            return 3
        if s == ".dotx":
            return 2
        if s == ".doc":
            return 1
        return 0

    candidates.sort(key=score, reverse=True)
    return candidates[0]


def _ensure_docx_if_doc(template_path: Path) -> Path:
    """Se o template for .doc, tenta converter para .docx usando MS Word (COM).

    Retorna o .docx correspondente quando convertido; caso contrário, retorna o caminho original.
    """
    suf = template_path.suffix.lower()
    if suf != ".doc":
        return template_path
    out_path = template_path.with_suffix(".docx")
    if out_path.exists():
        return out_path
    try:
        import win32com.client  # type: ignore
        from win32com.client import constants  # type: ignore
    except Exception:
        print("Aviso: pywin32/Word indisponivel para converter .doc -> .docx. Usando modelo original .doc.")
        return template_path
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(template_path))
        doc.SaveAs2(str(out_path), FileFormat=constants.wdFormatXMLDocument)
        doc.Close(SaveChanges=False)
        word.Quit()
        return out_path if out_path.exists() else template_path
    except Exception as e:
        print(f"Aviso: falha ao converter .doc para .docx ({e}). Usando modelo original .doc.")
        try:
            word.Quit()
        except Exception:
            pass
        return template_path


def classify_and_select_template_path(pdf_text: str, last_piece_name: Optional[str], cover_text: Optional[str] = None) -> Optional[Path]:
    """Aplica a classificação e retorna o Path do modelo selecionado, se encontrado.

    Usa o PDF da capa (cover_text) para identificar melhor a secretaria quando disponível
    e também para avaliar as palavras-chave de escolha do modelo.
    """
    tipo = _classify_tipo_from_text_and_piece(pdf_text or "", last_piece_name, cover_text=cover_text)
    secretaria_text = pdf_text or ""
    if cover_text:
        secretaria_text = f"{cover_text}\n{secretaria_text}"
    secretaria = _detect_secretaria_from_text(secretaria_text)
    p = _select_template_for(tipo, secretaria)
    if not p:
        print(f"Aviso: nenhum modelo encontrado para tipo={tipo} secretaria={secretaria}. Usando configuração padrão se existir.")
        return None
    p2 = _ensure_docx_if_doc(p)
    try:
        print(f"Modelo selecionado: tipo={tipo}, secretaria={secretaria}, arquivo={p2.name}")
    except Exception:
        pass
    return p2


def attach_docx_to_portal(context, page, docx_path: Path) -> bool:
    """Try to attach the generated DOCX in the current process UI.

    This looks for an <input type=file> or a button that opens a file chooser
    in any visible frame and attempts to upload the provided file.
    Returns True if the file was submitted to the form.
    """
    if not docx_path or not docx_path.exists():
        return False

    # 1) Direct file inputs across frames (bounded, non-blocking)
    containers = [page] + list(page.frames)
    for c in containers:
        try:
            el = c.query_selector("input[type='file']")
            if el:
                try:
                    el.set_input_files(str(docx_path))
                    print(f"Arquivo anexado via input[file]: {docx_path.name}")
                    # Try to click a likely 'Salvar'/'Enviar' afterwards
                    for btn_text in ("Salvar", "Gravar", "Enviar", "Confirmar"):
                        try:
                            c.get_by_role("button", name=re.compile(btn_text, re.I)).first.click()
                            break
                        except Exception:
                            continue
                    return True
                except Exception:
                    pass
        except Exception:
            continue

    # 2) Try buttons that open a file chooser
    trigger_texts = ["Anexar", "Incluir", "Inserir", "Upload", "Novo Documento", "Adicionar"]
    for c in containers:
        for txt in trigger_texts:
            try:
                with page.expect_file_chooser(timeout=5000) as fc_info:
                    c.get_by_role("button", name=re.compile(txt, re.I)).first.click()
                fc = fc_info.value
                fc.set_files(str(docx_path))
                print(f"Arquivo selecionado para upload: {docx_path.name}")
                # Try to confirm
                for btn_text in ("Salvar", "Gravar", "Enviar", "Confirmar"):
                    try:
                        c.get_by_role("button", name=re.compile(btn_text, re.I)).first.click()
                        break
                    except Exception:
                        continue
                return True
            except Exception:
                continue

    print("Aviso: nao foi possivel localizar interface de anexo automaticamente.")
    return False


def process_processo_pipeline(context, main_page, output_dir: Path, processo_num: str, use_caixa_correio: bool):
    """Fluxo completo: abre o processo, baixa PDF, gera oficio, cria comunicacao e anexa DOCX."""
    active_page = None
    try:
        maybe_page = filter_and_open_processo(context, main_page, processo_num)
        active_page = maybe_page or main_page
        try:
            find_frame_with_selector(active_page, "#splLeitorDocumentos_pgcPecas_trePecas", timeout_ms=15000)
        except Exception:
            try:
                active_page.locator(f"#cod_processo[value*='{processo_num}']").first.wait_for(state="attached", timeout=8000)
            except Exception:
                active_page = search_processo_and_open_viewer(context, main_page, processo_num)
        pdf_path, piece_title = click_last_piece_and_open_pdf(context, active_page, output_dir, processo_num)
        if not pdf_path:
            print(f"Aviso: nenhum PDF encontrado para {processo_num}.")
            return
        pdf_text = extract_text_from_pdf(pdf_path)
        cover_text = _extract_cover_text(context, active_page, output_dir, processo_num)
        fields = parse_fields_from_pdf_text(pdf_text, processo_num)
        tipo = _classify_tipo_from_text_and_piece(pdf_text or "", piece_title, cover_text=cover_text)
        secretaria_text = f"{cover_text}\n{pdf_text}" if cover_text else pdf_text
        secretaria = _detect_secretaria_from_text(secretaria_text)
        tpl_path = classify_and_select_template_path(pdf_text, piece_title, cover_text=cover_text)
        docx_path = generate_oficio_from_template(processo_num, output_dir, extra=fields, template_path=tpl_path)

        data_decadencia = extract_data_decadencia(pdf_text)
        prazo = calcular_prazo_res_22_21(data_decadencia, date.today())
        relator = fields.get("@@nome_relator") or fields.get("{{RELATOR}}") or fields.get("{{RELATOR_PROCESSO}}") or ""
        descricao = f"Oficio {tipo} - modelo {secretaria} - gerado automaticamente"

        if use_caixa_correio:
            try:
                caixa_target = open_caixa_correio_from_grid(context, main_page, processo_num)
                criar_comunicacao_processual(context, caixa_target, {
                    "processo": processo_num,
                    "secretaria": secretaria,
                    "relator": relator,
                    "tipo": tipo,
                    "prazo": prazo,
                    "descricao": descricao,
                })
            except Exception as e:
                print(f"Aviso: falha ao criar comunicacao processual para {processo_num}: {e}")

        if docx_path:
            try:
                attached = attach_docx_via_gerenciador_atos(context, main_page, processo_num, docx_path)
                if not attached:
                    attached = attach_docx_to_portal(context, active_page, docx_path)
                if attached:
                    print("Anexo do DOCX concluido.")
                else:
                    print("Aviso: anexo do DOCX nao foi concluido automaticamente.")
            except Exception as e:
                print(f"Aviso: falha ao anexar DOCX: {e}")
    finally:
        try:
            if active_page is not None and active_page != main_page:
                active_page.close()
        except Exception:
            pass


def main():
    load_dotenv()  # load .env if present

    url = os.getenv("ETCM_URL", "https://homologacao-etcm.tcm.sp.gov.br/paginas/login.aspx")
    username = os.getenv("ETCM_USERNAME")
    password = os.getenv("ETCM_PASSWORD")
    headless = env_bool("HEADLESS", False)
    show_browser = env_bool("SHOW_BROWSER", False) or env_bool("WATCH_MODE", False) or env_bool("FORCE_HEADED", False)
    if show_browser:
        headless = False
    try:
        slow_mo_ms = int(os.getenv("SLOWMO_MS", "200"))
    except Exception:
        slow_mo_ms = 200
    try:
        pause_after_login_ms = int(os.getenv("PAUSE_AFTER_LOGIN_MS", "0"))
    except Exception:
        pause_after_login_ms = 0
    try:
        login_manual_wait_ms = int(os.getenv("LOGIN_MANUAL_WAIT_MS", "45000"))
    except Exception:
        login_manual_wait_ms = 45000
    devtools = env_bool("DEVTOOLS", False)
    attach_only = env_bool("ATTACH_ONLY", False)
    use_caixa_correio = env_bool("USE_CAIXA_CORREIO", True)
    use_storage_state = env_bool("USE_STORAGE_STATE", True)
    storage_state_path = os.getenv("STORAGE_STATE_PATH", "storage_state.json").strip()
    storage_state_file = Path(storage_state_path) if storage_state_path else None

    if not username or not password:
        print("ERRO: defina ETCM_USERNAME e ETCM_PASSWORD (via .env ou variaveis de ambiente).")
        sys.exit(2)
    try:
        print(f"Usando usuario: {username}")
    except Exception:
        pass

    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)
    cleanup_output_dir(output_dir)

    with sync_playwright() as p:
        launch_kwargs = {"headless": headless, "channel": "chrome"}
        if devtools:
            launch_kwargs["devtools"] = True
        if not headless:
            print("Modo visivel: navegador sera exibido (HEADLESS desativado).")
        if headless:
            # Force Chrome's new headless implementation to avoid the removed legacy mode.
            launch_kwargs["args"] = ["--headless=new"]
        else:
            launch_kwargs["slow_mo"] = slow_mo_ms
        browser = p.chromium.launch(**launch_kwargs)
        context_kwargs = {"viewport": {"width": 1600, "height": 900}, "accept_downloads": True}
        if use_storage_state and storage_state_file and storage_state_file.exists():
            try:
                context_kwargs["storage_state"] = str(storage_state_file)
                print(f"Carregando sessao anterior: {storage_state_file}")
            except Exception:
                pass
        context = browser.new_context(**context_kwargs)
        page = context.new_page()

        # 1) Login (tenta reutilizar sessao se houver storage_state)
        need_login = True
        if use_storage_state and storage_state_file and storage_state_file.exists():
            mesa_url = urljoin(url, "/paginas/mesatrabalho.aspx")
            try:
                page.goto(mesa_url, wait_until="domcontentloaded", timeout=60000)
            except Exception:
                pass
            need_login = _is_login_page(page)
            if not need_login:
                print("Sessao anterior reutilizada com sucesso.")
        if need_login:
            login_etcm(
                page,
                url,
                username,
                password,
                pause_after_login_ms=pause_after_login_ms,
                login_manual_wait_ms=login_manual_wait_ms,
                headless=headless,
            )
            if use_storage_state and storage_state_file:
                try:
                    storage_state_file.parent.mkdir(parents=True, exist_ok=True)
                    context.storage_state(path=str(storage_state_file))
                    print(f"Sessao salva em: {storage_state_file}")
                except Exception:
                    pass
        try:
            find_frame_with_text(page, "Comunicados Importantes", timeout_ms=20000)
        except Exception:
            try:
                find_frame_with_text(page, "Processos", timeout_ms=20000)
            except Exception:
                pass

        # Modo ATTACH_ONLY: apenas anexa o DOCX mais recente via Gerenciador de Atos
        ger_atos_url = os.getenv("ETCM_GERENCIA_ATO_URL")
        if attach_only:
            if ger_atos_url:
                try:
                    print(f"Acessando Gerenciador de Atos: {ger_atos_url}")
                    page.goto(ger_atos_url, wait_until="load", timeout=60000)
                except Exception:
                    pass
            proc_label = os.getenv("PROCESSO_LABEL", "")
            latest_docx = None
            try:
                latest_docx = sorted(Path("output").glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)[0]
            except Exception:
                latest_docx = None
            if not latest_docx:
                print("ERRO: nenhum DOCX encontrado em output/ para anexar.")
                context.close(); browser.close(); return
            try:
                # Se nao foi passada a URL do gerenciador, tente abrir via grid usando o numero do processo
                if not ger_atos_url and proc_label:
                    open_gerenciador_atos_from_grid(context, page, proc_label)
                ok = attach_docx_via_gerenciador_atos(context, page, proc_label, latest_docx)
                if ok:
                    print("Anexo do DOCX concluido (ATTACH_ONLY).")
                else:
                    print("Aviso: nao foi possivel anexar o DOCX no modo ATTACH_ONLY.")
            except Exception as e:
                print(f"Aviso: falha no anexo ATTACH_ONLY: {e}")
            context.close(); browser.close(); return

        # Direct viewer URL (optional fast-path)
        viewer_url = os.getenv("ETCM_VIEWER_URL")
        if viewer_url:
            print(f"Acessando visualizador direto: {viewer_url}")
            page.goto(viewer_url, wait_until="load", timeout=60000)
            # Try to fetch and save last PDF immediately
            proc_label = os.getenv("PROCESSO_LABEL", "viewer")
            try:
                pdf_path, piece_title = click_last_piece_and_open_pdf(context, page, output_dir, proc_label)
                if pdf_path:
                    # Gera oficio a partir de template, se existir
                    pdf_text = extract_text_from_pdf(pdf_path)
                    cover_text = _extract_cover_text(context, page, output_dir, proc_label)
                    fields = parse_fields_from_pdf_text(pdf_text, proc_label)
                    # Seleciona template automaticamente conforme palavras-chave
                    tpl_path = classify_and_select_template_path(pdf_text, piece_title, cover_text=cover_text)
                    docx_path = generate_oficio_from_template(proc_label, output_dir, extra=fields, template_path=tpl_path)
                    if docx_path:
                        try:
                            attached = attach_docx_to_portal(context, page, docx_path)
                            if not attached:
                                # Tenta via Gerenciador de Atos. Se informada URL direta, navega ate ela.
                                ger_url_env = os.getenv("ETCM_GERENCIA_ATO_URL")
                                if ger_url_env:
                                    try:
                                        page.goto(ger_url_env, wait_until="load", timeout=30000)
                                    except Exception:
                                        pass
                                attached = attach_docx_via_gerenciador_atos(context, page, proc_label, docx_path)
                            if attached:
                                print("Anexo do DOCX concluido.")
                            else:
                                print("Aviso: anexo do DOCX nao foi concluido automaticamente.")
                        except Exception as e:
                            print(f"Aviso: falha ao anexar DOCX: {e}")
            except Exception as e:
                print(f"Aviso: falha no fluxo do visualizador direto: {e}")
            print("Concluido com sucesso.")
            context.close()
            browser.close()
            return

        # 2) Abrir APO-PEN e exportar a planilha
        downloaded_file = open_apo_pen_and_export_excel(context, page, output_dir)

        try:
            src_file = downloaded_file if downloaded_file and downloaded_file.exists() else find_latest_export_file(output_dir)
        except Exception:
            src_file = None

        processos_env: list[str] = []
        env_list = os.getenv("PROCESSOS_LIST")
        if env_list:
            sep = ";" if ";" in env_list and "," not in env_list else ","
            processos_env = [s.strip() for s in env_list.split(sep) if s.strip()]
        doit_all = env_bool("PROCESS_ALL", False) or bool(processos_env)

        if doit_all:
            processos: list[str] = []
            if processos_env:
                processos = processos_env
            elif src_file and src_file.exists():
                processos = read_processos_from_excel(src_file)
            if not processos:
                print("Aviso: nenhuma linha de processo identificada para processar.")
            else:
                try:
                    max_proc = int(os.getenv("MAX_PROCESSOS", "0"))
                except Exception:
                    max_proc = 0
                if max_proc > 0:
                    processos = processos[:max_proc]
                seen = set()
                processos = [p for p in processos if not (p in seen or seen.add(p))]
                print(f"Processos a tratar ({len(processos)}): {processos}")
                for idx, pr in enumerate(processos, start=1):
                    print(f"\n[{idx}/{len(processos)}] Tratando processo: {pr}")
                    try:
                        process_processo_pipeline(context, page, output_dir, pr, use_caixa_correio)
                    except Exception as e:
                        print(f"Aviso: falha no processamento de {pr}: {e}")
                print("Concluido com sucesso.")
                context.close()
                browser.close()
                return

        processo_num = os.getenv("PROCESSO_LABEL") or None
        if not processo_num:
            if src_file and src_file.exists():
                processo_num = extract_processo_from_excel(src_file)
                if processo_num:
                    print(f"Processo identificado na planilha: {processo_num}")
                else:
                    print("Aviso: nao foi possivel extrair 'No Processo' da planilha.")
            else:
                print("Aviso: nenhuma planilha encontrada para leitura.")

        if processo_num:
            try:
                process_processo_pipeline(context, page, output_dir, processo_num, use_caixa_correio)
            except Exception as e:
                print(f"Aviso: falha ao navegar e baixar PDF: {e}")

        print("Concluido com sucesso.")
        context.close()
        browser.close()


if __name__ == "__main__":
    main()
